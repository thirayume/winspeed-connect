# -*- coding: utf-8 -*-
"""WorldFert doc-gen: manifest-driven DOCX matrix (SRS, Tech Spec, User Guides) with
ISO doc-control + revision history + embedded mermaid diagrams. Source of truth = markdown."""
import os, re, glob, json, datetime, mistune
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ENT  = os.path.abspath(os.path.join(HERE, "..", ".."))
DIAG = os.path.join(ENT, "pipeline", "diagrams")

# ---------- manifest-driven (ห้าม hardcode โครงเอกสารในไฟล์นี้) ----------
MANIFEST_PATH = os.environ.get("WF_DOCX_MANIFEST") or os.path.join(HERE, "docx-manifest.json")
with open(MANIFEST_PATH, encoding="utf-8") as fh:
    MF = json.load(fh)

OUT  = os.path.join(ENT, *MF.get("outputDir", "01-DOCX/generated").split("/"))
FONT = MF.get("font", "Prompt")

_TH_MONTHS = ["มกราคม","กุมภาพันธ์","มีนาคม","เมษายน","พฤษภาคม","มิถุนายน",
              "กรกฎาคม","สิงหาคม","กันยายน","ตุลาคม","พฤศจิกายน","ธันวาคม"]
def _today_label():
    """docDate ว่างใน manifest = ใช้วันที่ปัจจุบันอัตโนมัติ (พ.ศ. + ค.ศ.)"""
    d = datetime.date.today()
    return (f"{d.day} {_TH_MONTHS[d.month-1]} {d.year+543} "
            f"({d.day} {d.strftime('%B')} {d.year})")

TODAY = MF.get("docDate") or _today_label()

def _app_version():
    """อ่านเลขเวอร์ชันจาก package.json ของ repo — ห้าม hardcode ในไฟล์นี้
    ชื่อไฟล์และหน้าปกจะได้ตรงกับ runtime เสมอโดยไม่ต้องแก้โค้ด"""
    for rel in (("..", "..", "package.json"),
                ("..", "..", "WSSale-App", "package.json")):
        candidate = os.path.abspath(os.path.join(ENT, *rel))
        if os.path.exists(candidate):
            try:
                with open(candidate, encoding="utf-8") as fh:
                    version = json.load(fh).get("version")
                if version:
                    return version
            except Exception:
                pass
    return MF.get("appVersion", "0.0.0")

APP_VERSION = MF.get("appVersion") or _app_version()
DOC_VERSION = MF.get("docVersion") or ("v" + APP_VERSION)
md = mistune.create_markdown(renderer=None, plugins=["table","strikethrough","url"])

M     = MF["documents"]
ROLES = MF["roles"]

# ---------- docx helpers ----------
def base(doc):
    st=doc.styles["Normal"]; st.font.name=FONT; st.font.size=Pt(10.5)
    st.element.rPr.rFonts.set(qn("w:eastAsia"),FONT); st.element.rPr.rFonts.set(qn("w:cs"),FONT)
def cs(r):
    r.font.name=FONT; rr=r._element.rPr
    if rr is not None and rr.rFonts is not None: rr.rFonts.set(qn("w:cs"),FONT); rr.rFonts.set(qn("w:eastAsia"),FONT)
def inline(par,nodes,b=False,i=False):
    for n in nodes or []:
        t=n.get("type")
        if t=="text": r=par.add_run(n.get("raw","")); r.bold=b; r.italic=i; cs(r)
        elif t=="strong": inline(par,n["children"],True,i)
        elif t=="emphasis": inline(par,n["children"],b,True)
        elif t=="codespan": r=par.add_run(n.get("raw","")); r.font.name="Consolas"; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xC7,0x25,0x4E); cs(r)
        elif t=="image":
            # ภาพใน markdown เคยถูกทิ้งทั้งหมดเพราะไม่มีกรณีนี้ ทำให้เอกสารไม่มีภาพหน้าจอเลย
            PENDING_IMAGES.append((n.get("url") or n.get("attrs",{}).get("url",""),
                                   "".join(c.get("raw","") for c in (n.get("children") or []))))
        elif t in ("link","strikethrough"): inline(par,n["children"],b,i)
        elif t in ("linebreak","softbreak"): par.add_run(" ")
        elif "children" in n: inline(par,n["children"],b,i)
        elif "raw" in n: r=par.add_run(n["raw"]); cs(r)

PENDING_IMAGES=[]
CURRENT_DIR=[None]

def flush_images(doc):
    """แทรกภาพที่พบใน markdown ลงเอกสาร พร้อมคำบรรยายใต้ภาพ

    ต้องทำหลังปิดย่อหน้าปัจจุบัน เพราะ python-docx แทรกรูปกลางย่อหน้าไม่ได้
    ภาพชิดซ้ายตามที่เจ้าของระบบขอ"""
    while PENDING_IMAGES:
        url,alt=PENDING_IMAGES.pop(0)
        if not url or url.startswith(("http://","https://")): continue
        base=CURRENT_DIR[0] or ENT
        path=os.path.normpath(os.path.join(base,url))
        if not os.path.exists(path):
            path=os.path.normpath(os.path.join(ENT,url.lstrip("./")))
        if not os.path.exists(path): continue
        try:
            doc.add_picture(path,width=Inches(6.2))
        except Exception as error:
            print("  !! แทรกภาพไม่ได้ %s (%s)"%(url,error)); continue
        doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.LEFT
        if alt:
            cap=doc.add_paragraph(); cap.alignment=WD_ALIGN_PARAGRAPH.LEFT
            r=cap.add_run(alt); r.italic=True; r.font.size=Pt(9)
            r.font.color.rgb=RGBColor(0x60,0x60,0x60); cs(r)
def code(doc,txt):
    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.2)
    r=p.add_run(txt.rstrip("\n")); r.font.name="Consolas"; r.font.size=Pt(8.5)
    sh=OxmlElement("w:shd"); sh.set(qn("w:fill"),"F2F2F2"); p._p.get_or_add_pPr().append(sh)
def table(doc,node):
    head=node["children"][0]; body=node["children"][1] if len(node["children"])>1 else None
    hc=head["children"]; nc=len(hc); rows=body["children"] if body else []
    tb=doc.add_table(rows=1,cols=nc); tb.style="Light Grid Accent 1"
    for j,c in enumerate(hc): tb.rows[0].cells[j].paragraphs[0].text=""; inline(tb.rows[0].cells[j].paragraphs[0],c["children"],True)
    for row in rows:
        cells=row["children"]; tr=tb.add_row().cells
        for j in range(nc):
            tr[j].paragraphs[0].text=""
            if j<len(cells): inline(tr[j].paragraphs[0],cells[j]["children"])
    doc.add_paragraph()
def rlist(doc,node,lvl=0):
    ordered=node.get("attrs",{}).get("ordered",False)
    for it in node["children"]:
        first=True
        for ch in it["children"]:
            ct=ch.get("type")
            if ct in ("block_text","paragraph"):
                p=doc.add_paragraph(style=("List Number" if ordered else "List Bullet") if first else None)
                if not first: p.paragraph_format.left_indent=Inches(0.5*(lvl+1))
                inline(p,ch["children"]); first=False
            elif ct=="list": rlist(doc,ch,lvl+1)
            elif ct=="block_code": code(doc,ch.get("raw",""))
def render(doc,tokens):
    for n in tokens:
        t=n.get("type")
        if t=="heading":
            lv=n.get("attrs",{}).get("level",1); p=doc.add_heading(level=min(lv,4)); inline(p,n["children"])
            for r in p.runs: cs(r)
        elif t=="paragraph":
            par=doc.add_paragraph(); par.alignment=WD_ALIGN_PARAGRAPH.LEFT
            inline(par,n["children"]); flush_images(doc)
            if not par.runs and not par.text.strip():
                par._p.getparent().remove(par._p)
        elif t=="block_code": code(doc,n.get("raw",""))
        elif t=="block_quote":
            for ch in n["children"]:
                if ch.get("type")=="paragraph":
                    p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.3); inline(p,ch["children"])
                    for r in p.runs: r.italic=True; r.font.color.rgb=RGBColor(0x55,0x55,0x55)
                else: render(doc,[ch])
        elif t=="list": rlist(doc,n)
        elif t=="table": table(doc,n)
        elif t=="thematic_break": doc.add_paragraph("_"*36)
        elif "children" in n: render(doc,n["children"])
def slice_section(txt,anchor):
    """ตัดเฉพาะหัวข้อที่ระบุ ไปจนถึงหัวข้อระดับเดียวกันหรือสูงกว่าอันถัดไป

    ใช้ให้คู่มือแต่ละบทบาทหยิบเฉพาะส่วนของตัวเองจากไฟล์รวมได้
    โดยไม่ต้องแตกไฟล์ใหม่และไม่ต้องคัดลอกเนื้อหาซ้ำ"""
    lines=txt.split("\n"); key=anchor.strip().lower()
    start=None; level=0
    for i,line in enumerate(lines):
        m=re.match(r"^(#{1,6})\s+(.*)$",line)
        if not m: continue
        if start is None:
            if key in m.group(2).strip().lower():
                start=i; level=len(m.group(1))
        elif len(m.group(1))<=level:
            return "\n".join(lines[start:i])
    return None if start is None else "\n".join(lines[start:])

def add_md(doc,path):
    ref,_,anchor=path.partition("#")
    fp=os.path.join(ENT,ref)
    if not os.path.exists(fp): return False
    txt=open(fp,encoding="utf-8").read().lstrip("﻿")
    if anchor:
        part=slice_section(txt,anchor)
        if part is None:
            print("  !! ไม่พบหัวข้อ '%s' ใน %s"%(anchor,ref)); return False
        txt=part
    CURRENT_DIR[0]=os.path.dirname(fp)
    render(doc,md(txt)); flush_images(doc); CURRENT_DIR[0]=None
    return True
def add_diagram(doc,stem,cap):
    png=os.path.join(DIAG,stem+".png")
    if not os.path.exists(png): return
    doc.add_page_break()
    h=doc.add_heading("แผนภาพ: "+cap,level=2)
    for r in h.runs: cs(r)
    doc.add_picture(png,width=Inches(6.6))
    doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER

DIAG_TITLES={"01-architecture":"สถาปัตยกรรมระบบ 3 ชั้น","02-so-lifecycle":"วงจรใบสั่งขาย (SO Lifecycle)","03-rebate-coupon-flow":"เส้นทางรีเบทและคูปอง","04-rebate-sequence":"ลำดับการทำงานของรีเบท","05-erd":"ผังความสัมพันธ์ข้อมูล (ERD)","06-rbac":"สิทธิ์ตามบทบาท (RBAC)","07-swimlane-order-to-cash":"ผังงานข้ามหน่วยงาน Order-to-Cash","08-uml-rebate-domain":"UML โดเมนรีเบท","09-current-system-context":"บริบทระบบปัจจุบัน","10-current-api-surface":"ขอบเขต API ปัจจุบัน","11-document-evidence-flow":"เส้นทางหลักฐานเอกสาร","13-runtime-architecture":"สถาปัตยกรรมขณะทำงาน","14-so-lifecycle-source-aligned":"วงจรใบสั่งขาย (ตรงกับโค้ด)","15-operational-erd":"ERD เชิงปฏิบัติการ","16-rbac-source-aligned":"RBAC (ตรงกับโค้ด)","17-order-to-cash-workflow":"กระบวนการ Order-to-Cash","18-rebate-domain-uml":"UML โดเมนรีเบท (ตรงกับโค้ด)","19-truckscale-as-is":"AS-IS · ก่อนเชื่อมสองทาง","20-truckscale-to-be":"TO-BE · หลังเชื่อมสองทาง","21-truckscale-ship-sequence":"ลำดับการชั่งออกและเขียนกลับ"}

def cover(doc,spec):
    for _ in range(3): doc.add_paragraph()
    t=doc.add_paragraph(); t.alignment=1; r=t.add_run("World Fert · WS-Sale-App"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=RGBColor(0x0C,0x44,0x7C); cs(r)
    t=doc.add_paragraph(); t.alignment=1; r=t.add_run(spec["title"]); r.bold=True; r.font.size=Pt(28); cs(r)
    for line in [f"Version {DOC_VERSION} · App {APP_VERSION} · {spec['scope']}", TODAY, f"Audience: {spec['audience']}", "Confidential — Client / Authorized Partner Use Only"]:
        p=doc.add_paragraph(); p.alignment=1; r=p.add_run(line); r.font.size=Pt(12); cs(r)
    doc.add_page_break()
    # doc control
    h=doc.add_heading("การควบคุมเอกสาร (Document Control)",level=1); [cs(r) for r in h.runs]
    tb=doc.add_table(rows=0,cols=2); tb.style="Light List Accent 1"
    for k,v in [("Document ID",spec["docid"]),("Product","WS-Sale-App"),("Client","World Fert Co., Ltd."),("Version",DOC_VERSION),("App build",APP_VERSION),("Date",TODAY),("Scope",spec["scope"]),("Audience",spec["audience"]),("Status","Released"),("Classification","Confidential")]:
        c=tb.add_row().cells; c[0].paragraphs[0].add_run(k).bold=True; c[1].paragraphs[0].add_run(v)
        for cell in c:
            for rr in cell.paragraphs[0].runs: cs(rr)
    doc.add_paragraph()
    h=doc.add_heading("ประวัติการแก้ไข (Revision History)",level=2); [cs(r) for r in h.runs]
    rt=doc.add_table(rows=1,cols=4); rt.style="Light Grid Accent 1"
    for j,x in enumerate(["Version","Date","Author","Change"]): rt.rows[0].cells[j].paragraphs[0].add_run(x).bold=True
    c=rt.add_row().cells; [c[j].paragraphs[0].add_run(v) for j,v in enumerate([DOC_VERSION,TODAY,"Doc-gen pipeline",f"Generated from source {APP_VERSION}"])]
    doc.add_paragraph()
    h=doc.add_heading("สารบัญ (Table of Contents)",level=1); [cs(r) for r in h.runs]
    p=doc.add_paragraph(); run=p.add_run(); fld=OxmlElement("w:fldSimple"); fld.set(qn("w:instr"),r'TOC \o "1-2" \h \z \u'); run._r.addnext(fld)
    doc.add_paragraph("(คลิกขวา → Update Field เพื่อสร้างเลขหน้า)")
    doc.add_page_break()

def _resolve_id(v):
    return str(v).replace("{VERSION}", DOC_VERSION)

def build(spec):
    doc=Document(); base(doc); cover(doc,spec)
    if spec.get("roles_summary"):
        h=doc.add_heading("สรุปการใช้งานตาม Role",level=1); [cs(r) for r in h.runs]
        for rn,ri in ROLES.items():
            hh=doc.add_heading(rn,level=2); [cs(r) for r in hh.runs]
            p=doc.add_paragraph(); p.add_run("เมนูที่เกี่ยวข้อง: ").bold=True; p.add_run(", ".join(ri["groups"]))
            for r in p.runs: cs(r)
        for s in (spec.get("sections") or MF.get("roleBaseSections", [])): add_md(doc,s)
    else:
        for s in spec["sections"]:
            add_md(doc,s); doc.add_page_break()
    for d in spec.get("diagrams",[]): add_diagram(doc,d,DIAG_TITLES.get(d,d))
    out=os.path.join(OUT,_resolve_id(spec["id"])+".docx"); doc.save(out); return out

def build_role(role,info):
    # rbac = รหัสบทบาทจริงในระบบ ต้องใช้รหัสนี้ในเอกสาร ไม่ใช่คีย์ใน manifest
    # ผู้ใช้เห็นรหัสนี้บนหน้าจอและใน audit log จึงต้องตรงกันเป๊ะ
    rbac=info.get("rbac",role); th=info.get("titleTh",role)
    spec={"id":f"User-Guide-{role}-{DOC_VERSION}","title":f"คู่มือผู้ใช้ตามบทบาท — {th}","docid":info["docid"],
          "audience":f"บทบาท {th} ({rbac})","scope":"Per-Role"}
    doc=Document(); base(doc); cover(doc,spec)
    h=doc.add_heading(f"หน้าที่และเมนูของบทบาท {th} ({rbac})",level=1); [cs(r) for r in h.runs]
    p=doc.add_paragraph(); p.add_run("กลุ่มเมนูที่ใช้: ").bold=True; p.add_run(", ".join(info["groups"]))
    for r in p.runs: cs(r)
    p=doc.add_paragraph(); p.add_run("รหัสบทบาทในระบบ: ").bold=True; p.add_run(rbac)
    p.add_run("  ·  ใช้รหัสนี้อ้างอิงเวลาแจ้งปัญหาและตรวจ audit log")
    for r in p.runs: cs(r)
    doc.add_paragraph("รายละเอียดหน้าจอดูใน Test Catalog / User Guide ด้านล่าง (เน้นเฉพาะส่วนที่บทบาทนี้ใช้).")
    for s in (info.get("sections") or MF.get("roleBaseSections", [])): add_md(doc,s)
    for d in info["diagrams"]: add_diagram(doc,d,DIAG_TITLES.get(d,d))
    out=os.path.join(OUT,_resolve_id(spec["id"])+".docx"); doc.save(out); return out

os.makedirs(OUT,exist_ok=True)
built=[]
for spec in M: built.append(build(spec))
for role,info in ROLES.items(): built.append(build_role(role,info))
print(f"built {len(built)} DOCX -> {OUT}")
for b in built: print("  -", os.path.basename(b))
