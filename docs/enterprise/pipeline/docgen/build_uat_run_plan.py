#!/usr/bin/env python3
"""Build the controlled UAT Full Loop Run Plan candidate from uat-cases.json.

The output deliberately remains Review/Candidate. It never promotes approval or
release status; those decisions require the named business and operational owners.
"""

from __future__ import annotations

import argparse
import hashlib
import json
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


NAVY = "12304A"
BLUE = "176B87"
TEAL = "2A9D8F"
SKY = "DCEEF5"
PALE = "F4F8FA"
LINE = "C9D7DF"
INK = "1B2633"
MUTED = "5C6B76"
WHITE = "FFFFFF"
GREEN = "DFF3E4"
AMBER = "FFF1CC"
RED = "FBE0DF"

HERE = Path(__file__).resolve().parent
ENTERPRISE = HERE.parent.parent
REPO = ENTERPRISE.parent.parent
DEFAULT_DATA = HERE / "uat-cases.json"
DEFAULT_OUTPUT = ENTERPRISE / "01-DOCX" / "candidate" / "UAT-Full-Loop-Run-Plan-v1.0.docx"


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, name="Prompt", size=10, bold=None, color=INK, italic=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)
    return run


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_cell_text(cell, text, *, bold=False, color=INK, size=9, font="Prompt", align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(str(text or ""))
    set_run_font(r, font, size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_field(paragraph, instruction: str):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, "Prompt", 8, color=MUTED)


def keep_with_next(paragraph):
    paragraph.paragraph_format.keep_with_next = True


def add_label(doc, text, fill=SKY, color=NAVY):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_text(cell, text, bold=True, color=color, size=11, font="Kanit")
    table.rows[0].height = Cm(0.75)
    return table


def add_note(doc, title, body, fill=PALE, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(0.35)
    table.columns[1].width = Cm(16.3)
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    table.cell(0, 0).text = ""
    p = table.cell(0, 1).paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{title}  "), "Kanit", 10, True, NAVY)
    set_run_font(p.add_run(body), "Prompt", 9, False, INK)
    set_cell_margins(table.cell(0, 1), top=120, bottom=120, start=140, end=140)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_meta_table(doc, rows, widths=(4.4, 12.4)):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].width = Cm(widths[0])
        cells[1].width = Cm(widths[1])
        shade(cells[0], BLUE)
        shade(cells[1], PALE)
        set_cell_text(cells[0], label, bold=True, color=WHITE, size=9)
        set_cell_text(cells[1], value, size=9)
        cant_split(table.rows[-1])
    return table


def add_grid_table(doc, headers, rows, widths=None, header_fill=BLUE, font_size=8.3):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for idx, h in enumerate(headers):
        shade(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], h, bold=True, color=WHITE, size=8.5)
        if widths:
            table.rows[0].cells[idx].width = Cm(widths[idx])
    set_repeat_table_header(table.rows[0])
    for r_idx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            shade(cells[idx], WHITE if r_idx % 2 else PALE)
            set_cell_text(cells[idx], value, size=font_size)
            if widths:
                cells[idx].width = Cm(widths[idx])
        cant_split(table.rows[-1])
    return table


def add_bullets(doc, items, level=0):
    for item in items:
        p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(str(item)), "Prompt", 9.5, color=INK)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        set_run_font(p.add_run(str(item)), "Prompt", 9.5, color=INK)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.add_run(text)
    keep_with_next(p)
    return p


def add_picture(doc, image_path: Path, caption: str, checks: list[str]):
    if not image_path.exists():
        add_note(doc, "ภาพประกอบยังไม่พร้อม", f"ไม่พบ {image_path.name}; ต้อง rerun/capture ก่อน release", RED, "A32828")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.keep_with_next = True
    inline = p.add_run().add_picture(str(image_path), width=Cm(16.7))
    inline._inline.docPr.set("descr", caption)
    inline._inline.docPr.set("title", f"UI illustration — {caption}")
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(3)
    set_run_font(cap.add_run(caption), "Prompt", 8.2, False, MUTED, True)
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(table.cell(0, 0), TEAL)
    shade(table.cell(0, 1), PALE)
    set_cell_text(table.cell(0, 0), "ตรวจในภาพ", bold=True, color=WHITE, size=9, font="Kanit")
    table.cell(0, 1).text = ""
    for idx, check in enumerate(checks, start=1):
        p2 = table.cell(0, 1).add_paragraph() if idx > 1 else table.cell(0, 1).paragraphs[0]
        p2.paragraph_format.space_after = Pt(1)
        set_run_font(p2.add_run(f"{idx}. {check}"), "Prompt", 8.7, color=INK)
    set_cell_margins(table.cell(0, 1), top=90, bottom=90, start=120, end=120)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Prompt"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Prompt")
    normal.font.size = Pt(10)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(4)

    for style_name, size, color, before, after in (
        ("Title", 30, WHITE, 0, 6),
        ("Heading 1", 18, NAVY, 14, 6),
        ("Heading 2", 14, BLUE, 10, 4),
        ("Heading 3", 11.5, TEAL, 7, 3),
    ):
        style = styles[style_name]
        style.font.name = "Kanit"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Kanit")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ("List Bullet", "List Bullet 2", "List Number"):
        style = styles[style_name]
        style.font.name = "Prompt"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Prompt")
        style.font.size = Pt(9.5)


def configure_page(doc: Document):
    section = doc.sections[0]
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)
    section.header_distance = Cm(0.6)
    section.footer_distance = Cm(0.65)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(17.7))
    table.columns[0].width = Cm(11.7)
    table.columns[1].width = Cm(6.0)
    shade(table.cell(0, 0), NAVY)
    shade(table.cell(0, 1), BLUE)
    set_cell_text(table.cell(0, 0), "WORLD FERT • WS-SALE-APP", bold=True, color=WHITE, size=8.5, font="Kanit")
    set_cell_text(table.cell(0, 1), "CONTROLLED CANDIDATE", bold=True, color=WHITE, size=8.2, font="Kanit", align=WD_ALIGN_PARAGRAPH.RIGHT)

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("WF-QA-022  |  Review v1.0-candidate  |  Confidential  |  Page "), "Prompt", 8, color=MUTED)
    add_field(p, "PAGE")
    set_run_font(p.add_run(" of "), "Prompt", 8, color=MUTED)
    add_field(p, "NUMPAGES")


def cover(doc, payload, source_commit, generated_at, e2e_status):
    top = doc.add_table(rows=1, cols=1)
    top.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(top.cell(0, 0), NAVY)
    cell = top.cell(0, 0)
    set_cell_margins(cell, top=700, bottom=700, start=500, end=500)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run("WORLD FERT\n"), "Kanit", 13, True, "8ED1DC")
    set_run_font(p.add_run("UAT FULL LOOP\nRUN PLAN"), "Kanit", 30, True, WHITE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(8)
    set_run_font(p2.add_run("WS-Sale-App • Sales → Approval → Scale → Warehouse → Reconciliation"), "Prompt", 11, color=WHITE)

    doc.add_paragraph()
    add_meta_table(doc, [
        ("Document ID", "WF-QA-022"),
        ("Version / Status", "v1.0-candidate / Review"),
        ("Product / Runtime", f"{payload.get('product', 'WS-Sale-App')} / {payload.get('runtimeVersion', '1.0.1')}"),
        ("Client", payload.get("client", "World Fert Co., Ltd.")),
        ("Source commit", source_commit),
        ("Generated", generated_at),
        ("Automated evidence", e2e_status),
        ("Classification", "Confidential — Client / Authorized Partner Use Only"),
    ])
    doc.add_paragraph()
    add_note(doc, "Purpose", "ใช้ควบคุมการทดสอบข้ามบทบาท บันทึกหลักฐาน และตัดสิน Go/No-Go; เหมาะทั้ง QA, business process owners และผู้ตรวจเอกสาร", SKY, BLUE)
    add_note(doc, "Release warning", "เอกสารนี้ไม่ใช่ UAT approval. ภาพ E2E เป็น illustrative evidence และต้อง rerun บน source/test hash ปัจจุบันก่อนลงนาม", AMBER, "C98500")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    set_run_font(p.add_run("Prepared for controlled review • 23 July 2026"), "Prompt", 9, color=MUTED)
    doc.add_page_break()


def build_document(data_path: Path, output: Path, source_commit: str, generated_at: str, e2e_status: str):
    payload = json.loads(data_path.read_text(encoding="utf-8"))
    cases = payload["cases"]
    doc = Document()
    configure_styles(doc)
    configure_page(doc)
    doc.core_properties.title = "UAT Full Loop Run Plan — WS-Sale-App"
    doc.core_properties.subject = "Controlled UAT plan and master test script"
    doc.core_properties.author = "World Fert / Solution Delivery Team"
    doc.core_properties.comments = "Generated from uat-cases.json; Review candidate only"

    cover(doc, payload, source_commit, generated_at, e2e_status)

    add_heading(doc, "Document control", 1)
    add_meta_table(doc, [
        ("Owner", "QA Lead / Business Process Owners"),
        ("Reviewers", "Solution Architect, Sales, Manager/Admin, Counter/Weighbridge, Warehouse, Accounting/WINSpeed, IT/Security"),
        ("Approver", "Business Sponsor / Authorized Go-Live Decision Owner"),
        ("Source of test content", "docs/enterprise/pipeline/docgen/uat-cases.json"),
        ("Related source", "UAT-FULL-LOOP-RUN-PLAN.md, TEST-STRATEGY.md, UAT-AND-SIGNOFF.md, RTM"),
        ("Generated artefact", "01-DOCX/candidate/UAT-Full-Loop-Run-Plan-v1.0.docx"),
    ])

    add_heading(doc, "Revision and approval history", 2)
    add_grid_table(doc, ["Version", "Date", "Change", "Prepared", "Reviewed", "Approved"], [
        ["v1.0-candidate", "23 Jul 2026", "Initial source-driven UAT candidate with 18 cases and 5-role illustrated walkthrough", "Generator / QA", "", ""],
    ], [2.4, 2.3, 6.2, 2.2, 2.2, 2.2])

    add_heading(doc, "Approval record", 2)
    add_grid_table(doc, ["Role", "Name", "Decision", "Date", "Evidence / Signature"], [
        [r, "", "Pending", "", ""] for r in [
            "Sales Process Owner", "Manager/Admin Owner", "Counter/Weighbridge Owner", "Warehouse Owner",
            "Accounting/WINSpeed Owner", "IT/Security/Operations", "QA Lead", "Business Sponsor",
        ]
    ], [4.3, 3.0, 2.5, 2.5, 4.5])
    doc.add_page_break()

    add_heading(doc, "Contents and navigation", 1)
    add_grid_table(doc, ["Section", "What the reader gets", "Primary users"], [
        ["1. Executive run guide", "Scope, business loop and what automation does not prove", "All participants"],
        ["2. Entry / exit / decision", "Readiness checklist and Go/No-Go rule", "QA, owners, sponsor"],
        ["3. Run control / evidence", "Run metadata, evidence naming and control", "QA, IT"],
        ["4. Illustrated full loop", "Five-role steps, current UI illustrations and checks", "End users, trainers, QA"],
        ["5. Manual gates", "TruckScale, WINSpeed, rebate, BCP, security and performance", "Business/IT owners"],
        ["6–7. Schedule / defects", "Coordination, severity, retest and escalation", "QA, delivery team"],
        ["8. Master test cases", "Detailed 18-case script generated from the shared JSON source", "Testers, reviewers"],
        ["9. Final decision", "Gate results, conditions, signatures and decision evidence", "Decision authority"],
    ], [4.0, 8.6, 4.2])
    add_note(doc, "Navigation note", "หัวข้อทุกระดับใช้ Word heading styles; เปิด Navigation Pane เพื่อกระโดดไปยัง section ได้ทันที", PALE, TEAL)

    add_heading(doc, "1. Executive run guide", 1)
    doc.add_paragraph("เอกสารนี้พาผู้ทดสอบจบหนึ่ง business loop และต่อด้วย manual gates ที่ automation ยังแทนไม่ได้ เพื่อให้ผลทดสอบอธิบายได้ทั้งกับผู้ใช้และทีม technical")
    add_label(doc, "SALES  →  MANAGER / ADMIN  →  COUNTER SALES  →  WAREHOUSE  →  ADMIN / ACCOUNTING", SKY, NAVY)
    add_bullets(doc, [
        "ใช้ Run ID และ Test Data ID เดียวกันตลอด loop เพื่อ trace รายการได้",
        "หยุดที่ hold point เมื่อผล critical เป็น Fail/Blocked; ห้ามข้ามเพื่อให้จบ script",
        "ทุก Pass/Fail ต้องมีหลักฐานและผู้ตรวจ; Automated evidence ต้อง hash ตรงกับ source ปัจจุบัน",
        "TruckScale online, WINSpeed reconciliation, security, performance และ restore ต้องมี manual acceptance",
    ])

    add_heading(doc, "1.1 Scope and objectives", 2)
    add_grid_table(doc, ["In scope", "Objective / acceptance focus"], [
        ["Sales order and verification", "ข้อมูลลูกค้า/สินค้า/จำนวน/ทะเบียน trace ได้ และ confirm ไม่ข้าม verification"],
        ["Approval and audit", "actor/time/reason/status ถูกต้องและตรวจย้อนหลังได้"],
        ["TruckScale and counter", "health/freshness/source ชัดเจน; degraded fallback ปลอดภัย"],
        ["Warehouse execution", "pick/load/ship ตามลำดับ ปริมาณไม่เกิน และ actor trace ได้"],
        ["Paper/accounting/reconciliation", "SHIPPED state, paper exceptions, WINSpeed references และยอดตรวจเทียบได้"],
        ["Operational assurance", "RBAC, performance, backup/restore, BCP และ evidence control"],
    ], [5.2, 11.6])

    add_heading(doc, "1.2 What this plan does not prove automatically", 2)
    add_bullets(doc, [
        "ไม่ได้รับรอง production approval, ISO certification หรือ security compliance โดยอัตโนมัติ",
        "ไม่แทน business owner signatures, accounting reconciliation หรือ live TruckScale test",
        "ไม่อนุญาตให้ใช้ผล E2E เก่าหลัง source/test hash เปลี่ยน",
    ])

    add_heading(doc, "2. Entry, exit and decision criteria", 1)
    add_heading(doc, "2.1 Entry checklist", 2)
    add_grid_table(doc, ["ID", "Control", "Owner", "Status", "Evidence"], [
        ["ENT-01", "Build, commit, source/test hash recorded and matched", "QA", "Not checked", ""],
        ["ENT-02", "Frontend/API/SQL health confirmed; dependency state recorded", "IT", "Not checked", ""],
        ["ENT-03", "Named test users and least-privilege roles ready", "Admin/Security", "Not checked", ""],
        ["ENT-04", "Synthetic/masked test data and reset plan ready", "Data Owner", "Not checked", ""],
        ["ENT-05", "Evidence root, defect tracker and communication channel ready", "QA", "Not checked", ""],
        ["ENT-06", "Backup/rollback/restore owner and prerequisites ready", "IT", "Not checked", ""],
    ], [1.6, 7.0, 3.0, 2.3, 3.2])

    add_heading(doc, "2.2 Exit / Go-No-Go rule", 2)
    add_grid_table(doc, ["Decision", "Rule"], [
        ["Accept / Go", "Critical 100% Pass; overall ≥90%; no open Sev-1/Sev-2; reconciliation, security, restore, evidence and signatures complete"],
        ["Conditional Go", "Critical all pass; only low-risk defects with owner/due/mitigation and signed residual-risk acceptance"],
        ["Reject / No-Go", "Critical workflow/security/accounting integrity fails, rollback/restore unavailable, or evidence cannot be traced"],
    ], [3.5, 13.3])

    add_heading(doc, "3. Run control and evidence", 1)
    add_meta_table(doc, [
        ("Run ID", ""), ("Environment / URL", ""), ("App build / source commit", f"{payload.get('runtimeVersion', '1.0.1')} / {source_commit}"),
        ("Source hash / test hash", ""), ("Test data batch", ""), ("Start / End", ""),
        ("QA Lead / Business Owner", ""), ("Evidence root / defect tracker", ""),
        ("Dependency health", "Frontend: __  API: __  SQL: __  TruckScale: __  WINSpeed: __"),
    ])
    add_note(doc, "Current evidence status", e2e_status, AMBER, "C98500")
    add_heading(doc, "3.1 Evidence naming and minimum metadata", 2)
    add_bullets(doc, [
        "ชื่อไฟล์: <RunID>_<CaseID>_<Step>_<Actor>_<YYYYMMDD-HHmm>.<ext>",
        "บันทึก actor, role, environment, build/commit, timestamp, expected/actual และ reviewer",
        "ไฟล์ API/log/query ต้อง mask token, password, secret และ PII ที่ไม่จำเป็น",
        "Evidence register ต้องมี controlled location และ SHA-256 ก่อน sign-off",
    ])

    screenshots_root = REPO / "test-results" / "artifacts"
    ui_root = ENTERPRISE / "05-UI-SCREENSHOTS"
    shots = {
        "UAT-E2E-002": ui_root / "02-sales-pos.png",
        "UAT-E2E-003": ui_root / "18-approval-policy.png",
        "UAT-E2E-004": next(iter(screenshots_root.glob("*29cb0*/*.png")), Path("missing-scale.png")),
        "UAT-E2E-005": ui_root / "04-warehouse.png",
        "UAT-E2E-006": ui_root / "05-paper-trail.png",
    }

    add_heading(doc, "4. Illustrated five-role full loop", 1)
    add_note(doc, "ภาพประกอบ", "ใช้ current UI captures ที่อ่านง่ายสำหรับ Sales, approval policy, Warehouse และ Paper Trail; ภาพ TruckScale มาจาก automated evidence. ทุกภาพเป็น instructional illustration ไม่ใช่ acceptance evidence จนกว่าจะ rerun/capture พร้อม Run ID และ source hash ปัจจุบัน", AMBER, "C98500")
    primary_ids = ["UAT-E2E-002", "UAT-E2E-003", "UAT-E2E-004", "UAT-E2E-005", "UAT-E2E-006"]
    check_map = {
        "UAT-E2E-002": ["ชื่อผู้ใช้/role เป็น SALES", "SO/test reference trace ได้", "สถานะเริ่มต้นและข้อมูลหลักตรง test data"],
        "UAT-E2E-003": ["ผู้ตรวจเป็น MANAGER/ADMIN ตามลำดับ", "verification/confirm status ถูกต้อง", "audit actor/time แสดงได้"],
        "UAT-E2E-004": ["health state ของ TruckScale ชัดเจน", "freshness/source ไม่ทำให้เข้าใจผิด", "มี degraded/fallback guidance"],
        "UAT-E2E-005": ["ผู้ใช้ WAREHOUSE และรายการเดียวกับต้นทาง", "pick/load/ship state ตามลำดับ", "quantity/vehicle/document prerequisite ถูกต้อง"],
        "UAT-E2E-006": ["final status SHIPPED", "Paper Trail แสดง actor/time/reference", "พร้อมส่งต่อ accounting reconciliation"],
    }
    case_by_id = {c["id"]: c for c in cases}
    for seq, case_id in enumerate(primary_ids, start=1):
        case = case_by_id[case_id]
        add_heading(doc, f"4.{seq} {case['titleTh']}", 2)
        add_meta_table(doc, [
            ("Case / Actor", f"{case_id} / {case['actor']}"),
            ("Preconditions", case["preconditions"]),
            ("Test data", case["testData"]),
            ("Expected", case["expected"]),
        ])
        add_heading(doc, "Operator steps", 3)
        add_numbered(doc, case["steps"])
        add_picture(doc, shots[case_id], f"Current UI illustration — {case_id} / {shots[case_id].name}; capture with Run ID required before acceptance", check_map[case_id])
        add_note(doc, "Evidence to attach", case["evidence"], PALE, TEAL)

    add_heading(doc, "5. Manual acceptance gates", 1)
    manual_cases = [c for c in cases if c["type"] == "Manual"]
    add_grid_table(doc, ["Case", "Area / objective", "Actor", "Critical", "Evidence"], [
        [c["id"], f"{c['area']} — {c['titleTh']}", c["actor"], "Yes" if c["critical"] else "No", c["evidence"]]
        for c in manual_cases
    ], [2.3, 6.4, 3.0, 1.6, 3.7], font_size=7.8)
    add_note(doc, "Mandatory hold", "ห้ามลงนาม Accept หาก UAT-MAN-001, 002, 006 หรือ 007 ยัง Not Run/Blocked เว้นแต่ final decision owner ระบุ No-Go หรือทำ formal risk acceptance ที่อนุญาตโดย policy", RED, "A32828")

    add_heading(doc, "6. Suggested schedule and coordination", 1)
    add_grid_table(doc, ["Window", "Activity", "Owner", "Hold point"], [
        ["T-2 days", "Baseline/config/data/evidence preparation", "QA + IT", "Entry review"],
        ["T-1 day", "Automated regression/E2E on frozen candidate", "QA", "Hash/evidence review"],
        ["Day 1 AM", "Sales + approval + negative verification gate", "Sales/Manager", "SO control review"],
        ["Day 1 PM", "Counter/TruckScale + Warehouse", "Factory/Warehouse", "Scale/ship review"],
        ["Day 2 AM", "Paper, WINSpeed, rebate/CN, giveaway", "Admin/Accounting", "Reconciliation review"],
        ["Day 2 PM", "RBAC, performance, restore/BCP", "IT/Security/QA", "Operational review"],
        ["Day 3", "Defect retest and final sign-off", "All owners", "Go/No-Go"],
    ], [2.4, 7.1, 3.3, 4.0])

    add_heading(doc, "7. Defect, retest and escalation", 1)
    add_grid_table(doc, ["Severity", "Definition", "Response"], [
        ["Sev-1", "Data loss/corruption, security breach, process cannot continue", "Stop run; No-Go; fix and impacted full regression"],
        ["Sev-2", "Critical control/integration wrong; workaround unsafe", "Fix before sign-off; targeted + end-to-end retest"],
        ["Sev-3", "Function wrong with safe workaround", "Owner/due; Conditional only with signed risk acceptance"],
        ["Sev-4", "Cosmetic/documentation/usability minor", "Backlog allowed with evidence and owner"],
    ], [2.2, 7.1, 7.5])
    add_bullets(doc, [
        "ทุก Fail ต้องมี defect ID, actual result, evidence, owner และ target date",
        "ทุก Closed defect ต้องมี retest result/evidence และ reviewer",
        "การแก้ critical workflow ต้อง rerun full loop ที่ได้รับผลกระทบ ไม่ทดสอบเฉพาะปุ่มที่แก้",
    ])

    add_heading(doc, "8. Detailed master test cases", 1)
    add_note(doc, "Source-driven appendix", f"รายละเอียด {len(cases)} cases ต่อไปนี้สร้างจาก uat-cases.json เดียวกับ workbook เพื่อลดความขัดแย้งระหว่าง artefacts", SKY, BLUE)
    for idx, case in enumerate(cases, start=1):
        add_heading(doc, f"8.{idx} {case['id']} — {case['titleTh']}", 2)
        add_meta_table(doc, [
            ("Type / Area", f"{case['type']} / {case['area']}"),
            ("Actor / Critical", f"{case['actor']} / {'Yes' if case['critical'] else 'No'}"),
            ("Requirements", ", ".join(case.get("requirementIds", []))),
            ("Automation", case.get("automationRef") or "Manual"),
            ("Initial status", case.get("status", "Not Run")),
        ])
        add_heading(doc, "Preconditions", 3)
        doc.add_paragraph(case["preconditions"])
        add_heading(doc, "Test data", 3)
        doc.add_paragraph(case["testData"])
        add_heading(doc, "Steps", 3)
        add_numbered(doc, case["steps"])
        add_heading(doc, "Expected result and evidence", 3)
        add_grid_table(doc, ["Expected result", "Required evidence", "Actual / Status / Defect"], [[case["expected"], case["evidence"], ""]], [6.2, 5.7, 4.9], font_size=8.3)

    add_heading(doc, "9. Final decision record", 1)
    add_grid_table(doc, ["Gate", "Target", "Actual", "Evidence", "Decision"], [
        ["Critical cases", "100% Pass", "", "", "Pending"],
        ["Overall cases", ">=90% Pass", "", "", "Pending"],
        ["Open Sev-1/Sev-2", "0", "", "", "Pending"],
        ["TruckScale online", "Pass", "", "", "Pending"],
        ["WINSpeed reconciliation", "Pass", "", "", "Pending"],
        ["Security/RBAC", "Pass", "", "", "Pending"],
        ["Restore/BCP", "Pass", "", "", "Pending"],
        ["Evidence/signatures", "Complete", "", "", "Pending"],
    ], [3.8, 2.6, 2.6, 5.1, 2.7])
    add_meta_table(doc, [
        ("Final decision", "Pending / Go / Conditional Go / No-Go"),
        ("Conditions / residual risks", ""),
        ("Rollback / contingency reference", ""),
        ("Decision meeting / evidence", ""),
        ("Authorized approver / date", ""),
    ])

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output)
    digest = hashlib.sha256(output.read_bytes()).hexdigest().upper()
    manifest = {
        "documentId": "WF-QA-022",
        "status": "Review",
        "output": str(output),
        "sha256": digest,
        "sourceData": str(data_path),
        "sourceCommit": source_commit,
        "generatedAt": generated_at,
        "caseCount": len(cases),
        "automatedCount": sum(1 for c in cases if c["type"] == "Automated"),
        "manualCount": sum(1 for c in cases if c["type"] == "Manual"),
        "e2eStatus": e2e_status,
    }
    manifest_path = output.with_suffix(output.suffix + ".manifest.json")
    manifest_path.write_text(json.dumps(manifest, ensure_ascii=False, indent=2), encoding="utf-8")
    return manifest


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--data", type=Path, default=DEFAULT_DATA)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--source-commit", default="UNVERIFIED")
    parser.add_argument("--generated-at", default=datetime.now(timezone.utc).astimezone().isoformat(timespec="seconds"))
    parser.add_argument("--e2e-status", default="RERUN REQUIRED — tracked source drift detected; previous screenshots are illustrative only")
    args = parser.parse_args()
    result = build_document(args.data.resolve(), args.output.resolve(), args.source_commit, args.generated_at, args.e2e_status)
    print(json.dumps(result, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
