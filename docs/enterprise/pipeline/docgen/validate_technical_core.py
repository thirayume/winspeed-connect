#!/usr/bin/env python3
"""Validate technical-core DOCX candidates and every rendered page."""

from __future__ import annotations

import argparse
import hashlib
import json
import re
import zipfile
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from docx import Document
from lxml import etree
from PIL import Image, ImageChops, ImageDraw, ImageFont, ImageStat


HERE = Path(__file__).resolve().parent
ENTERPRISE = HERE.parent.parent
REPO = ENTERPRISE.parent.parent
OUTPUT = ENTERPRISE / "01-DOCX" / "candidate" / "technical-core"
REPORT = ENTERPRISE / "pipeline" / "reports" / "technical-core-validation-report.json"

DOCS = [
    ("srs", "SRS-v1.0-candidate.docx"),
    ("analysis-design", "Analysis-and-Design-v1.0-candidate.docx"),
    ("technical-spec", "Technical-Specification-v1.0-candidate.docx"),
]

NS = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
}


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def add_check(checks: list[dict[str, Any]], check_id: str, status: str, detail: str, severity="high"):
    checks.append({"id": check_id, "status": status, "severity": severity, "detail": detail})


def numeric_page_key(path: Path) -> int:
    match = re.search(r"(\d+)$", path.stem)
    return int(match.group(1)) if match else 0


def page_metrics(path: Path) -> dict[str, Any]:
    with Image.open(path) as source:
        image = source.convert("RGB")
        width, height = image.size
        background = Image.new("RGB", image.size, (255, 255, 255))
        diff = ImageChops.difference(image, background).convert("L")
        threshold = diff.point(lambda p: 255 if p > 10 else 0)
        bbox = threshold.getbbox()
        histogram = threshold.histogram()
        ink_pixels = histogram[255]
        ink_ratio = ink_pixels / (width * height)
        edge = max(3, round(min(width, height) * 0.002))
        edge_regions = {
            "top": threshold.crop((0, 0, width, edge)),
            "bottom": threshold.crop((0, height - edge, width, height)),
            "left": threshold.crop((0, 0, edge, height)),
            "right": threshold.crop((width - edge, 0, width, height)),
        }
        edge_ink = {name: bool(region.getbbox()) for name, region in edge_regions.items()}
        return {
            "path": path.relative_to(REPO).as_posix(),
            "width": width,
            "height": height,
            "inkRatio": round(ink_ratio, 6),
            "contentBox": list(bbox) if bbox else None,
            "edgeInk": edge_ink,
            "blank": bbox is None or ink_ratio < 0.001,
        }


def make_contact_sheets(page_paths: list[Path], output_dir: Path) -> list[dict[str, Any]]:
    output_dir.mkdir(parents=True, exist_ok=True)
    sheets: list[dict[str, Any]] = []
    per_sheet = 6
    columns = 3
    thumb_width = 650
    label_height = 36
    gap = 24
    margin = 30
    for start in range(0, len(page_paths), per_sheet):
        batch = page_paths[start : start + per_sheet]
        thumbs: list[tuple[Image.Image, str]] = []
        for page in batch:
            with Image.open(page) as source:
                image = source.convert("RGB")
                ratio = thumb_width / image.width
                thumb = image.resize((thumb_width, round(image.height * ratio)), Image.Resampling.LANCZOS)
                thumbs.append((thumb, f"Page {numeric_page_key(page)}"))
        thumb_height = max(image.height for image, _ in thumbs)
        rows = (len(thumbs) + columns - 1) // columns
        sheet_width = margin * 2 + columns * thumb_width + (columns - 1) * gap
        sheet_height = margin * 2 + rows * (label_height + thumb_height) + (rows - 1) * gap
        sheet = Image.new("RGB", (sheet_width, sheet_height), (225, 231, 235))
        draw = ImageDraw.Draw(sheet)
        font = ImageFont.load_default(size=22)
        for index, (thumb, label) in enumerate(thumbs):
            row, col = divmod(index, columns)
            x = margin + col * (thumb_width + gap)
            y = margin + row * (label_height + thumb_height + gap)
            draw.text((x, y + 4), label, fill=(21, 51, 74), font=font)
            sheet.paste(thumb, (x, y + label_height))
        sheet_index = start // per_sheet + 1
        output_path = output_dir / f"contact-{sheet_index:02d}.png"
        sheet.save(output_path, optimize=True)
        sheets.append(
            {
                "path": output_path.relative_to(REPO).as_posix(),
                "pages": [numeric_page_key(path) for path in batch],
                "sha256": sha256(output_path),
                "width": sheet.width,
                "height": sheet.height,
            }
        )
    return sheets


def inspect_docx(path: Path) -> dict[str, Any]:
    doc = Document(path)
    headings = [p.text for p in doc.paragraphs if p.style and p.style.name.startswith("Heading")]
    body_text = "\n".join(p.text for p in doc.paragraphs)
    with zipfile.ZipFile(path) as archive:
        document_xml = archive.read("word/document.xml")
        styles_xml = archive.read("word/styles.xml")
        root = etree.fromstring(document_xml)
        tables = root.xpath("//w:tbl", namespaces=NS)
        multi_row = 0
        multi_row_with_header = 0
        fixed_layout = 0
        for table in tables:
            rows = table.xpath("./w:tr", namespaces=NS)
            if len(rows) >= 2:
                multi_row += 1
                if rows[0].xpath("./w:trPr/w:tblHeader", namespaces=NS):
                    multi_row_with_header += 1
            if table.xpath("./w:tblPr/w:tblLayout[@w:type='fixed']", namespaces=NS):
                fixed_layout += 1
        doc_prs = root.xpath("//wp:docPr", namespaces=NS)
        alt_text = [node.get("descr", "").strip() for node in doc_prs]
        fonts_present = {
            "Prompt": b"Prompt" in styles_xml or b"Prompt" in document_xml,
            "Kanit": b"Kanit" in styles_xml or b"Kanit" in document_xml,
        }
    return {
        "paragraphs": len(doc.paragraphs),
        "headings": len(headings),
        "headingSample": headings[:12],
        "tables": len(doc.tables),
        "fixedLayoutTables": fixed_layout,
        "multiRowTables": multi_row,
        "multiRowTablesWithHeaderSemantics": multi_row_with_header,
        "inlineShapes": len(doc.inline_shapes),
        "drawingObjects": len(doc_prs),
        "drawingObjectsWithAltText": sum(1 for value in alt_text if value),
        "fontsPresent": fonts_present,
        "reviewMarker": "REVIEW" in body_text.upper() or "Review" in body_text,
        "releasedMarkerInBody": bool(re.search(r"\bStatus\s*[:|-]?\s*Released\b", body_text, re.IGNORECASE)),
    }


def validate_one(key: str, filename: str, render_root: Path) -> dict[str, Any]:
    checks: list[dict[str, Any]] = []
    docx_path = OUTPUT / filename
    manifest_path = docx_path.with_suffix(docx_path.suffix + ".manifest.json")
    render_dir = render_root / key
    page_dir = render_dir / "word-pdfium"
    render_report_path = render_dir / "render-report.json"

    if not docx_path.exists():
        add_check(checks, "docx.exists", "FAIL", str(docx_path))
        return {"key": key, "status": "FAILED", "checks": checks}
    add_check(checks, "docx.exists", "PASS", f"{docx_path.name} · {docx_path.stat().st_size} bytes")

    manifest = json.loads(manifest_path.read_text(encoding="utf-8-sig"))
    actual_hash = sha256(docx_path)
    add_check(
        checks,
        "docx.manifestHash",
        "PASS" if manifest["sha256"] == actual_hash else "FAIL",
        f"manifest={manifest['sha256']} actual={actual_hash}",
    )
    add_check(
        checks,
        "docx.status",
        "PASS" if manifest.get("status") == "Review" else "FAIL",
        f"status={manifest.get('status')}",
    )

    source_mismatches = []
    for source in manifest.get("sources", []):
        path = REPO / source["path"]
        if not path.exists() or sha256(path) != source["sha256"]:
            source_mismatches.append(source["path"])
    add_check(
        checks,
        "source.hashes",
        "PASS" if not source_mismatches else "FAIL",
        "all source hashes match" if not source_mismatches else f"mismatch: {source_mismatches}",
    )

    structure = inspect_docx(docx_path)
    add_check(checks, "structure.headings", "PASS" if structure["headings"] >= 12 else "FAIL", f"{structure['headings']} headings")
    add_check(
        checks,
        "structure.fixedTables",
        "PASS" if structure["fixedLayoutTables"] == structure["tables"] else "FAIL",
        f"{structure['fixedLayoutTables']}/{structure['tables']} fixed-layout tables",
    )
    add_check(
        checks,
        "accessibility.tableHeaders",
        "PASS" if structure["multiRowTablesWithHeaderSemantics"] == structure["multiRowTables"] else "FAIL",
        f"{structure['multiRowTablesWithHeaderSemantics']}/{structure['multiRowTables']} multi-row tables have repeated header semantics",
    )
    add_check(
        checks,
        "accessibility.imageAltText",
        "PASS" if structure["drawingObjectsWithAltText"] == structure["drawingObjects"] else "FAIL",
        f"{structure['drawingObjectsWithAltText']}/{structure['drawingObjects']} drawing objects have alt text",
    )
    add_check(
        checks,
        "style.thaiFonts",
        "PASS" if all(structure["fontsPresent"].values()) else "FAIL",
        json.dumps(structure["fontsPresent"]),
    )
    add_check(checks, "status.reviewMarker", "PASS" if structure["reviewMarker"] else "FAIL", "Review marker present")

    render_report = json.loads(render_report_path.read_text(encoding="utf-8-sig"))
    page_paths = sorted(page_dir.glob("page-*.png"), key=numeric_page_key)
    add_check(
        checks,
        "render.pageCount",
        "PASS" if len(page_paths) == render_report["pages"] and page_paths else "FAIL",
        f"report={render_report['pages']} png={len(page_paths)} method={render_report['method']}",
    )
    metrics = [page_metrics(path) for path in page_paths]
    blanks = [index + 1 for index, item in enumerate(metrics) if item["blank"]]
    edge_touch = [
        index + 1
        for index, item in enumerate(metrics)
        if any(item["edgeInk"].values())
    ]
    dimensions = sorted({(item["width"], item["height"]) for item in metrics})
    add_check(checks, "render.noBlankPages", "PASS" if not blanks else "FAIL", f"blank pages={blanks}")
    add_check(
        checks,
        "render.noEdgeClipping",
        "PASS" if not edge_touch else "FAIL",
        "no content touches raster edge" if not edge_touch else f"edge ink pages={edge_touch}",
    )
    add_check(
        checks,
        "render.consistentDimensions",
        "PASS" if len(dimensions) == 1 else "FAIL",
        f"dimensions={dimensions}",
    )
    contacts = make_contact_sheets(page_paths, render_dir / "contact-sheets")
    add_check(
        checks,
        "render.contactSheets",
        "PASS" if contacts and sum(len(item["pages"]) for item in contacts) == len(page_paths) else "FAIL",
        f"{len(contacts)} sheets cover {sum(len(item['pages']) for item in contacts)} pages",
    )

    failed = [check for check in checks if check["status"] == "FAIL" and check["severity"] == "high"]
    return {
        "key": key,
        "status": "FAILED" if failed else "PASSED",
        "document": docx_path.relative_to(REPO).as_posix(),
        "sha256": actual_hash,
        "structure": structure,
        "render": {
            "method": render_report["method"],
            "pages": len(page_paths),
            "dimensions": dimensions,
            "pageMetrics": metrics,
            "contactSheets": contacts,
        },
        "checks": checks,
    }


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument(
        "--render-root",
        type=Path,
        default=OUTPUT / "qa-render-r3",
        help="Directory containing srs, analysis-design, and technical-spec render folders.",
    )
    args = parser.parse_args()
    render_root = args.render_root.resolve()
    documents = [validate_one(key, filename, render_root) for key, filename in DOCS]
    status = "PASSED" if all(item["status"] == "PASSED" for item in documents) else "FAILED"
    report = {
        "schemaVersion": 1,
        "kind": "WS-Sale-App Technical Core Validation",
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "status": status,
        "renderRoot": render_root.relative_to(REPO).as_posix(),
        "documents": documents,
        "summary": {
            "documents": len(documents),
            "pages": sum(item.get("render", {}).get("pages", 0) for item in documents),
            "checks": sum(len(item.get("checks", [])) for item in documents),
            "failedChecks": sum(
                1
                for item in documents
                for check in item.get("checks", [])
                if check["status"] == "FAIL"
            ),
        },
    }
    REPORT.parent.mkdir(parents=True, exist_ok=True)
    REPORT.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report["summary"] | {"status": status}, ensure_ascii=True, indent=2))
    return 0 if status == "PASSED" else 1


if __name__ == "__main__":
    raise SystemExit(main())
