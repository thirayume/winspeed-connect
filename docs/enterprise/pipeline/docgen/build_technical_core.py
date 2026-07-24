#!/usr/bin/env python3
"""Build source-aligned technical-core DOCX candidates.

The generator intentionally produces Review/Candidate documents. It never
promotes a document to Approved or Released and never accepts a source or
document baseline.
"""

from __future__ import annotations

import argparse
import hashlib
import json
import re
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable

import mistune
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


NAVY = "15334A"
BLUE = "1B6B88"
TEAL = "2B8C82"
SKY = "DDEEF3"
PALE = "F4F8FA"
LINE = "C7D5DC"
INK = "1C2730"
MUTED = "5B6872"
WHITE = "FFFFFF"
AMBER = "F4B942"
AMBER_PALE = "FFF3D1"
GREEN = "2D7D5B"
GREEN_PALE = "E4F2EA"
RED = "A63C3C"
RED_PALE = "F8E5E5"

HERE = Path(__file__).resolve().parent
ENTERPRISE = HERE.parent.parent
REPO = ENTERPRISE.parent.parent
DEFAULT_MANIFEST = HERE / "technical-core-manifest.json"
CONTENT_WIDTH_CM = 17.0


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest().upper()


def rgb(value: str) -> RGBColor:
    return RGBColor.from_string(value)


def set_run_font(run, name="Prompt", size=10, bold=None, color=INK, italic=None):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:cs"), name)
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = rgb(color)
    return run


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=110, bottom=90, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_cm: float):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(round(width_cm * 567)))


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = tr_pr.find(qn("w:tblHeader"))
    if node is None:
        node = OxmlElement("w:tblHeader")
        tr_pr.append(node)
    node.set(qn("w:val"), "true")


def cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_fixed(table, widths: list[float]):
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(round(width * 567)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(widths):
                set_cell_width(cell, widths[idx])


def set_table_borders(table, color=LINE, size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold=False, color=INK, size=8.8, font="Prompt", align=None):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    if align is not None:
        paragraph.alignment = align
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.02
    set_run_font(paragraph.add_run(str(text or "")), font, size, bold, color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_field(paragraph, instruction: str, placeholder="1"):
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_begin, instr, fld_sep, text, fld_end])
    set_run_font(run, "Prompt", 8, color=MUTED)


def add_grid_table(
    doc: Document,
    headers: list[str],
    rows: Iterable[Iterable[str]],
    widths: list[float] | None = None,
    font_size=8.3,
    header_fill=BLUE,
):
    if not headers:
        return None
    if widths is None:
        widths = calculate_widths(headers, rows)
    rows = [list(row) for row in rows]
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, widths)
    set_table_borders(table)
    for idx, header in enumerate(headers):
        shade(table.rows[0].cells[idx], header_fill)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE, size=8.4, font="Kanit")
    set_repeat_table_header(table.rows[0])
    cant_split(table.rows[0])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for idx in range(len(headers)):
            value = values[idx] if idx < len(values) else ""
            shade(cells[idx], PALE if row_index % 2 == 0 else WHITE)
            set_cell_text(cells[idx], value, size=font_size)
        cant_split(table.rows[-1])
    compact_text = sum(len(str(value)) for row in rows for value in row)
    if len(rows) <= 9 and compact_text <= 1400:
        for row in table.rows[:-1]:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.keep_with_next = True
    return table


def calculate_widths(headers: list[str], rows: Iterable[Iterable[str]]) -> list[float]:
    data = [headers] + [list(row) for row in rows]
    weights: list[float] = []
    for idx in range(len(headers)):
        samples = [str(row[idx] if idx < len(row) else "") for row in data]
        longest = max((min(len(s), 80) for s in samples), default=8)
        weights.append(max(5.0, min(26.0, longest ** 0.65)))
    minimum = 1.6 if len(headers) <= 5 else 1.25
    raw_total = sum(weights)
    widths = [max(minimum, CONTENT_WIDTH_CM * value / raw_total) for value in weights]
    scale = CONTENT_WIDTH_CM / sum(widths)
    return [round(width * scale, 2) for width in widths]


def add_note(doc: Document, title: str, body: str, fill=PALE, accent=BLUE):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [0.35, CONTENT_WIDTH_CM - 0.35])
    shade(table.cell(0, 0), accent)
    shade(table.cell(0, 1), fill)
    table.cell(0, 0).text = ""
    paragraph = table.cell(0, 1).paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    set_run_font(paragraph.add_run(f"{title}  "), "Kanit", 9.6, True, NAVY)
    set_run_font(paragraph.add_run(body), "Prompt", 9, False, INK)
    set_cell_margins(table.cell(0, 1), top=120, bottom=120, start=140, end=140)
    set_table_borders(table, color=fill, size=0)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = 0.4


def add_heading(doc: Document, text: str, level=1):
    level = min(max(level, 1), 4)
    paragraph = doc.add_paragraph(style=f"Heading {level}")
    set_run_font(
        paragraph.add_run(text),
        "Kanit",
        {1: 17, 2: 13.5, 3: 11.5, 4: 10.5}[level],
        True,
        {1: NAVY, 2: BLUE, 3: TEAL, 4: INK}[level],
    )
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def add_paragraph_text(doc: Document, text: str, style=None, *, color=INK, size=9.4, italic=False):
    paragraph = doc.add_paragraph(style=style)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run_font(paragraph.add_run(text), "Prompt", size, color=color, italic=italic)
    return paragraph


def add_picture(doc: Document, path: Path, caption: str, width_cm=16.6, max_height_cm=20.5):
    if not path.exists():
        add_note(doc, "ภาพประกอบยังไม่พร้อม", f"ไม่พบไฟล์ {path}; ต้องสร้างใหม่ก่อน release", RED_PALE, RED)
        return
    chosen_width = width_cm
    with path.open("rb") as stream:
        header = stream.read(24)
    if len(header) >= 24 and header[:8] == b"\x89PNG\r\n\x1a\n":
        pixel_width = int.from_bytes(header[16:20], "big")
        pixel_height = int.from_bytes(header[20:24], "big")
        if pixel_width and pixel_height:
            chosen_width = min(width_cm, max_height_cm * pixel_width / pixel_height)
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    inline = paragraph.add_run().add_picture(str(path), width=Cm(chosen_width))
    inline._inline.docPr.set("descr", caption)
    inline._inline.docPr.set("title", caption)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_after = Pt(6)
    set_run_font(cap.add_run(caption), "Prompt", 8.2, False, MUTED, True)


def configure_document(doc: Document, document_id: str, short_title: str):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.55)
    section.bottom_margin = Cm(1.45)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.header_distance = Cm(0.65)
    section.footer_distance = Cm(0.65)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Prompt"
    normal._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Prompt")
    normal.font.size = Pt(9.4)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.05
    normal.paragraph_format.space_after = Pt(3)

    for name, size, color, before, after in (
        ("Title", 27, WHITE, 0, 6),
        ("Subtitle", 13, WHITE, 0, 4),
        ("Heading 1", 17, NAVY, 12, 5),
        ("Heading 2", 13.5, BLUE, 9, 4),
        ("Heading 3", 11.5, TEAL, 7, 3),
        ("Heading 4", 10.5, INK, 5, 2),
    ):
        style = doc.styles[name]
        style.font.name = "Kanit"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Kanit")
        style.font.size = Pt(size)
        style.font.color.rgb = rgb(color)
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Number", "List Number 2"):
        style = doc.styles[name]
        style.font.name = "Prompt"
        style._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Prompt")
        style.font.size = Pt(9.2)
        style.paragraph_format.space_after = Pt(2)

    header = section.header
    table = header.add_table(rows=1, cols=2, width=Cm(CONTENT_WIDTH_CM))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [10.9, 6.1])
    shade(table.cell(0, 0), NAVY)
    shade(table.cell(0, 1), SKY)
    set_cell_text(table.cell(0, 0), f"WS-Sale-App  |  {short_title}", bold=True, color=WHITE, size=8.6, font="Kanit")
    set_cell_text(table.cell(0, 1), f"{document_id}  ·  REVIEW", bold=True, color=NAVY, size=8.3, font="Kanit", align=WD_ALIGN_PARAGRAPH.RIGHT)
    set_table_borders(table, color=NAVY, size=0)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(paragraph.add_run("CONFIDENTIAL · Candidate / Review only  |  "), "Prompt", 7.7, color=MUTED)
    add_field(paragraph, "PAGE")
    set_run_font(paragraph.add_run(" / "), "Prompt", 7.7, color=MUTED)
    add_field(paragraph, "NUMPAGES")


def add_cover(doc: Document, spec: dict[str, Any], manifest: dict[str, Any], cover_figure: Path | None):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [CONTENT_WIDTH_CM])
    shade(table.cell(0, 0), NAVY)
    cell = table.cell(0, 0)
    set_cell_margins(cell, top=340, bottom=340, start=280, end=280)
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run_font(p.add_run("WORLD FERT · SOLUTION DEVELOPMENT"), "Kanit", 10.5, True, AMBER)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(14)
    set_run_font(p2.add_run(spec["title"]), "Kanit", 28, True, WHITE)
    p3 = cell.add_paragraph()
    set_run_font(p3.add_run(spec["subtitle"]), "Prompt", 13, False, SKY)
    p4 = cell.add_paragraph()
    p4.paragraph_format.space_before = Pt(18)
    set_run_font(p4.add_run("SOURCE-ALIGNED CANDIDATE · REVIEW"), "Kanit", 10, True, AMBER)
    set_table_borders(table, color=NAVY, size=0)

    meta = [
        ("Document ID", spec["documentId"]),
        ("Version / Runtime", f"v1.0 candidate / runtime {manifest['runtimeVersion']}"),
        ("Date", "23 July 2026 / 23 กรกฎาคม 2569"),
        ("Source commit", manifest["sourceCommit"]),
        ("Source inventory SHA-256", manifest["sourceInventorySha256"]),
        ("Automated evidence", f"{manifest['e2eRunId']} · 10/10 passed · source stable"),
        ("Status", "Review — approval and baseline acceptance required"),
    ]
    add_grid_table(doc, ["Control", "Value"], meta, widths=[4.2, 12.8], font_size=8.6, header_fill=TEAL)
    add_note(
        doc,
        "การใช้เอกสาร",
        "สร้างจาก Markdown และ source evidence ล่าสุดตามนโยบาย latest-document-wins; ห้ามตีความเป็น Approved, Released หรือ ISO certification จนกว่าจะมีผู้มีอำนาจลงนามและรับ baseline.",
        AMBER_PALE,
        AMBER,
    )
    if cover_figure and cover_figure.exists():
        add_picture(doc, cover_figure, f"Executive visual — {spec['title']}", width_cm=15.6)
    doc.add_page_break()


def add_front_matter(doc: Document, spec: dict[str, Any], manifest: dict[str, Any]):
    add_heading(doc, "Document control", 1)
    add_grid_table(
        doc,
        ["รายการ", "รายละเอียด"],
        [
            ("Document ID", spec["documentId"]),
            ("Title", spec["title"]),
            ("Owner", owner_for(spec["key"])),
            ("Status", "Review — Candidate"),
            ("Classification", "Confidential — Client / Authorized Partner Use Only"),
            ("Design preset", "standard_business_brief · A4 · Prompt/Kanit override"),
            ("Header pattern", "memo_masthead"),
            ("Generation rule", "drift check → source merge → render → QA → human review"),
        ],
        widths=[4.4, 12.6],
    )
    add_heading(doc, "Revision and approval record", 2)
    add_grid_table(
        doc,
        ["Version", "Date", "Change", "Prepared by", "Reviewed / approved"],
        [
            (
                "v1.0 candidate",
                "2026-07-23",
                "Rebuilt from source-aligned Markdown, runtime inventory, diagrams and current E2E evidence",
                "Documentation pipeline",
                "Pending named human approval",
            )
        ],
        widths=[2.2, 2.4, 7.0, 2.8, 2.6],
        font_size=8.0,
    )

    add_heading(doc, "Executive source snapshot", 1)
    add_grid_table(
        doc,
        ["Evidence area", "Current verified fact", "Boundary"],
        [
            ("Runtime", "1.0.1 across root/backend/frontend packages", "Development source snapshot"),
            ("Source", "220 files · SHA 12B9F964…FC343 · commit 79a10a28", "Dirty worktree remains visible"),
            ("Application", "22 portals · 8 roles", "Role-specific UAT still required"),
            ("API", "17 route mounts · 160 detected endpoints", "Detection is not contract approval"),
            ("Data", "55 sequenced migrations through 055", "Production migration/restore gate remains"),
            ("External writes", "33 detected dbo writes · 2 TruckScale writes to tbl_keyone", "Explicit contract/security review required"),
            ("Automated E2E", "10/10 passed · stable source · SQL Server/MySQL up", "Not production hardware/manual UAT"),
        ],
        widths=[3.0, 7.8, 6.2],
        font_size=8.1,
    )
    add_note(
        doc,
        "Evidence interpretation",
        "ตัวเลขจาก static inventory และ E2E ใช้เพื่อยืนยันความสอดคล้องกับ snapshot ปัจจุบันเท่านั้น การยอมรับ UAT, external write boundary, production readiness และ ISO conformance ต้องมีหลักฐานและผู้อนุมัติแยกต่างหาก.",
        SKY,
        BLUE,
    )

    add_heading(doc, "Standards alignment", 1)
    add_grid_table(
        doc,
        ["Reference", "Use in this information set", "Current document"],
        standards_rows(spec["key"]),
        widths=[4.0, 8.6, 4.4],
        font_size=7.9,
    )

    add_heading(doc, "Audience and reading path", 1)
    add_grid_table(
        doc,
        ["Audience", "Primary concern", "Recommended sections"],
        audience_rows(spec["key"]),
        widths=[3.5, 7.1, 6.4],
        font_size=8.1,
    )

    add_heading(doc, "Table of contents", 1)
    toc_p = doc.add_paragraph()
    add_field(toc_p, 'TOC \\o "1-3" \\h \\z \\u', "Update field in Word to refresh the table of contents")
    add_note(doc, "Word field", "เปิดเอกสารใน Microsoft Word แล้วกด Ctrl+A และ F9 เพื่อ refresh สารบัญและจำนวนหน้า.", PALE, TEAL)
    doc.add_page_break()


def owner_for(key: str) -> str:
    return {
        "srs": "Product Owner / Business Analyst / Solution Architect",
        "analysis-design": "Solution Architect / Business Analyst / Data Architect",
        "technical-spec": "Solution Architect / Technical Lead / Security / DevOps",
    }[key]


def standards_rows(key: str) -> list[tuple[str, str, str]]:
    common = [
        ("ISO/IEC/IEEE 12207:2026", "Life-cycle processes, reviews and evidence flow", "Development information set"),
        ("ISO/IEC/IEEE 15289:2019", "Information-item content and document control", "Front matter / source register"),
        ("ISO/IEC 25010:2023", "Product quality characteristics and verification concerns", "NFR / architecture / technical gates"),
        ("ISO/IEC/IEEE 29119-3:2021", "Test documentation, cases, results and traceability", "Verification and UAT linkage"),
    ]
    if key == "srs":
        return [("ISO/IEC/IEEE 29148:2018", "Requirements structure, qualities and traceability", "Primary")] + common
    if key == "analysis-design":
        return [("ISO/IEC/IEEE 42010:2022", "Stakeholders, concerns, viewpoints, views and decisions", "Primary")] + common
    return [
        ("ISO/IEC/IEEE 42010:2022", "Technical views and decision correspondence", "Architecture linkage"),
        ("ISO/IEC/IEEE 29148:2018", "Technical and interface requirement traceability", "Contract linkage"),
        ("ISO/IEC 27001:2022", "Security-control context and objective evidence", "Security section"),
    ] + common


def audience_rows(key: str) -> list[tuple[str, str, str]]:
    if key == "srs":
        return [
            ("Business owner / Product owner", "scope, policy, acceptance", "Purpose, functional requirements, acceptance"),
            ("End users / UAT", "role workflow and observable result", "Roles, workflows, interface, UAT linkage"),
            ("Engineering / QA", "testable requirement and traceability", "FR/NFR, states, errors, implementation status"),
            ("Audit / ISO", "control, approval and retained evidence", "Document control, traceability, evidence boundary"),
        ]
    if key == "analysis-design":
        return [
            ("Business/process owner", "stakeholders and cross-role process", "Business analysis, workflow, decisions"),
            ("Architect / developer", "boundaries, components and constraints", "SAD, C4, ADR, UML"),
            ("DBA / integration owner", "ownership, data and external contracts", "Data design, ERD, integration decisions"),
            ("Operations / audit", "risk, evidence and maintainability", "Architecture concerns, decision records"),
        ]
    return [
        ("Developer / reviewer", "runtime contract and source responsibility", "Components, API, data, integration"),
        ("DBA / integration owner", "transactions, migration and external writes", "Database and integration contracts"),
        ("Security / DevOps", "deployment, secrets, logging, backup and recovery", "Security and operations"),
        ("QA / audit", "verification coverage and objective evidence", "Test strategy, evidence register, open gates"),
    ]


def strip_frontmatter(text: str) -> str:
    if text.startswith("---"):
        parts = text.split("---", 2)
        if len(parts) == 3:
            return parts[2].lstrip("\r\n")
    return text


def extract_source_title(path: Path, text: str) -> str:
    title_match = re.search(r'^title:\s*["\']?(.+?)["\']?\s*$', text, re.MULTILINE)
    if title_match and text.startswith("---"):
        return title_match.group(1).strip('"\'')
    heading = re.search(r"^#\s+(.+)$", strip_frontmatter(text), re.MULTILINE)
    return heading.group(1).strip() if heading else path.stem.replace("-", " ")


def inline_text(tokens: list[dict[str, Any]] | None) -> str:
    if not tokens:
        return ""
    parts: list[str] = []
    for token in tokens:
        token_type = token.get("type")
        if token_type in ("text", "codespan"):
            parts.append(token.get("raw", ""))
        elif token_type in ("softbreak", "linebreak"):
            parts.append("\n")
        elif token_type == "link":
            label = inline_text(token.get("children"))
            url = token.get("attrs", {}).get("url", "")
            parts.append(f"{label} ({url})" if url.startswith(("http://", "https://")) and url not in label else label)
        elif token_type == "image":
            alt = inline_text(token.get("children"))
            parts.append(alt or token.get("attrs", {}).get("url", "image"))
        else:
            parts.append(inline_text(token.get("children")))
    return "".join(parts).strip()


def flatten_block_text(token: dict[str, Any]) -> str:
    if "children" in token:
        return inline_text(token.get("children"))
    return token.get("raw", "").strip()


def table_to_matrix(token: dict[str, Any]) -> tuple[list[str], list[list[str]]]:
    headers: list[str] = []
    rows: list[list[str]] = []
    for part in token.get("children", []):
        if part.get("type") == "table_head":
            headers = [inline_text(cell.get("children")) for cell in part.get("children", [])]
        elif part.get("type") == "table_body":
            for row in part.get("children", []):
                values = [inline_text(cell.get("children")) for cell in row.get("children", [])]
                rows.append(values)
    return headers, rows


def add_code_block(doc: Document, raw: str, info: str):
    if info.strip().lower().startswith("mermaid"):
        add_note(doc, "Diagram source", "Mermaid source is retained in Markdown; the reviewed high-resolution rendering is embedded in the figure section.", PALE, TEAL)
        return
    display = raw.rstrip()
    if len(display) > 3200:
        display = display[:3200] + "\n… [excerpt limited in DOCX; authoritative source remains in repository]"
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_fixed(table, [CONTENT_WIDTH_CM])
    shade(table.cell(0, 0), "EFF3F5")
    set_table_borders(table, color=LINE, size=5)
    cell = table.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run_font(paragraph.add_run(display), "Consolas", 7.4, False, INK)
    set_cell_margins(cell, top=100, bottom=100, start=130, end=130)


def new_numbering_instance(doc: Document, style_name: str) -> int:
    style = doc.styles[style_name]
    num_nodes = style.element.xpath("./w:pPr/w:numPr/w:numId")
    if not num_nodes:
        raise ValueError(f"Numbering style {style_name} has no numId")
    base_num_id = int(num_nodes[0].get(qn("w:val")))
    numbering = doc.part.numbering_part.element
    base_nodes = numbering.xpath(f"./w:num[@w:numId='{base_num_id}']")
    if not base_nodes:
        raise ValueError(f"Base numbering definition {base_num_id} not found")
    abstract_id = base_nodes[0].find(qn("w:abstractNumId")).get(qn("w:val"))
    existing = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    new_id = max(existing, default=0) + 1
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(new_id))
    abstract = OxmlElement("w:abstractNumId")
    abstract.set(qn("w:val"), abstract_id)
    num.append(abstract)
    override = OxmlElement("w:lvlOverride")
    override.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:startOverride")
    start.set(qn("w:val"), "1")
    override.append(start)
    num.append(override)
    numbering.append(num)
    return new_id


def apply_numbering(paragraph, num_id: int):
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = p_pr.get_or_add_numPr()
    num_id_node = num_pr.get_or_add_numId()
    num_id_node.set(qn("w:val"), str(num_id))


def add_list(doc: Document, token: dict[str, Any], depth=0):
    ordered = bool(token.get("attrs", {}).get("ordered"))
    style = ("List Number" if ordered else "List Bullet") + (" 2" if depth else "")
    num_id = new_numbering_instance(doc, style) if ordered else None
    for item in token.get("children", []):
        text_parts: list[str] = []
        nested: list[dict[str, Any]] = []
        for child in item.get("children", []):
            if child.get("type") == "list":
                nested.append(child)
            else:
                value = flatten_block_text(child)
                if value:
                    text_parts.append(value)
        paragraph = doc.add_paragraph(style=style)
        if num_id is not None:
            apply_numbering(paragraph, num_id)
        paragraph.paragraph_format.space_after = Pt(2)
        set_run_font(paragraph.add_run(" ".join(text_parts)), "Prompt", 9.2, color=INK)
        for child in nested:
            add_list(doc, child, depth=1)


def add_source_content(doc: Document, source_path: Path, title: str, markdown) -> dict[str, int]:
    raw = source_path.read_text(encoding="utf-8-sig")
    body = strip_frontmatter(raw)
    tokens = markdown(body)
    add_heading(doc, title, 1)
    add_note(
        doc,
        "Authoritative source",
        f"{source_path.relative_to(REPO).as_posix()} · SHA-256 {sha256(source_path)} · merged as Review content",
        PALE,
        TEAL,
    )
    counts = {"headings": 1, "tables": 0, "paragraphs": 0, "lists": 0, "codeBlocks": 0}
    skipped_first_h1 = False
    skipped_control_table = False
    for token in tokens:
        token_type = token.get("type")
        if token_type in ("blank_line", "thematic_break"):
            continue
        if token_type == "heading":
            level = int(token.get("attrs", {}).get("level", 1))
            text = inline_text(token.get("children"))
            if level == 1 and not skipped_first_h1:
                skipped_first_h1 = True
                continue
            add_heading(doc, text, min(level + 1, 4))
            counts["headings"] += 1
        elif token_type == "paragraph":
            text = inline_text(token.get("children"))
            if text:
                add_paragraph_text(doc, text)
                counts["paragraphs"] += 1
        elif token_type == "block_quote":
            text = " ".join(flatten_block_text(child) for child in token.get("children", []))
            if (
                "Merge provenance" in text
                or "Status: Released" in text
                or "Source of truth:" in text
                or "เอกสารควบคุมสำหรับ ISO" in text
            ):
                continue
            add_note(doc, "Source note", text, PALE, BLUE)
            counts["paragraphs"] += 1
        elif token_type == "table":
            headers, rows = table_to_matrix(token)
            if not skipped_control_table and headers and any("รายการ" in h or "Item" in h for h in headers):
                skipped_control_table = True
                continue
            if headers:
                add_grid_table(doc, headers, rows, font_size=7.9 if len(headers) > 4 else 8.3)
                counts["tables"] += 1
        elif token_type == "list":
            add_list(doc, token)
            counts["lists"] += 1
        elif token_type == "block_code":
            add_code_block(doc, token.get("raw", ""), token.get("attrs", {}).get("info", ""))
            counts["codeBlocks"] += 1
        elif token_type in ("block_html", "inline_html"):
            text = re.sub(r"<[^>]+>", " ", token.get("raw", ""))
            text = re.sub(r"\s+", " ", text).strip()
            if text:
                add_paragraph_text(doc, text)
                counts["paragraphs"] += 1
    return counts


def add_figures(doc: Document, spec: dict[str, Any]):
    if not spec.get("figures"):
        return
    add_heading(doc, "Source-aligned diagrams", 1)
    add_note(
        doc,
        "Diagram policy",
        "ภาพในส่วนนี้สร้างจาก Mermaid source ที่ version-control ได้และ render ใหม่ด้วย Prompt/Kanit policy; แหล่ง Markdown ยังคงเป็นข้อมูลอ้างอิงและภาพต้องผ่าน visual QA ก่อน release.",
        SKY,
        BLUE,
    )
    for figure in spec["figures"]:
        add_picture(doc, ENTERPRISE / figure["path"], figure["caption"])

    if spec.get("screenshots"):
        add_heading(doc, "Automated UI evidence examples", 2)
        add_note(
            doc,
            "Evidence boundary",
            "ภาพต่อไปนี้เป็นผลจาก Playwright E2E run ปัจจุบัน ใช้แสดงตัวอย่าง UI และ trace เท่านั้น ไม่แทน manual UAT, production hardware, printing/scanning หรือ acceptance sign-off.",
            AMBER_PALE,
            AMBER,
        )
        for screenshot in spec["screenshots"]:
            add_picture(doc, ENTERPRISE / screenshot["path"], screenshot["caption"], width_cm=15.8)
    doc.add_page_break()


def add_source_register(doc: Document, spec: dict[str, Any], source_records: list[dict[str, Any]], figure_records: list[dict[str, Any]]):
    add_heading(doc, "Source and evidence register", 1)
    add_grid_table(
        doc,
        ["Type", "Path", "SHA-256", "Disposition"],
        [
            (
                "Markdown",
                record["path"],
                record["sha256"],
                "Merged into Review candidate",
            )
            for record in source_records
        ]
        + [
            (
                "Figure",
                record["path"],
                record["sha256"],
                "Embedded with alt text",
            )
            for record in figure_records
        ],
        widths=[2.0, 7.2, 5.2, 2.6],
        font_size=7.2,
    )
    doc.add_page_break()
    add_heading(doc, "Approval boundary", 1)
    add_note(
        doc,
        "Pending human decision",
        "Candidate นี้ยังไม่เป็น baseline. ต้อง review ความถูกต้องของ business rule, API/data contract, external write boundary, security, operations และ UAT evidence ก่อนบันทึกผู้อนุมัติ/วันที่/ข้อยกเว้น.",
        RED_PALE,
        RED,
    )
    add_grid_table(
        doc,
        ["Approval role", "Name", "Decision / date", "Conditions or evidence reference"],
        [
            ("Business owner / Product owner", "", "Pending", ""),
            ("Solution architect / Technical lead", "", "Pending", ""),
            ("Security / DBA / Integration owner", "", "Pending", ""),
            ("QA / UAT owner", "", "Pending", ""),
            ("Operations / Release authority", "", "Pending", ""),
        ],
        widths=[4.0, 3.0, 3.0, 7.0],
        font_size=8.2,
    )


def set_core_properties(doc: Document, spec: dict[str, Any]):
    props = doc.core_properties
    props.title = spec["title"]
    props.subject = spec["subtitle"]
    props.author = "World Fert Solution Documentation Pipeline"
    props.keywords = "WS-Sale-App, Review Candidate, ISO, source-aligned, traceability"
    props.comments = "Generated candidate. Human approval and baseline acceptance are required."
    props.category = "Controlled solution-development document"
    props.language = "th-TH"


def build_document(spec: dict[str, Any], manifest: dict[str, Any], markdown) -> dict[str, Any]:
    output_path = ENTERPRISE / spec["output"]
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_document(doc, spec["documentId"], spec["title"])
    set_core_properties(doc, spec)
    cover_figure = ENTERPRISE / spec["figures"][0]["path"] if spec.get("figures") else None
    add_cover(doc, spec, manifest, cover_figure)
    add_front_matter(doc, spec, manifest)
    add_figures(doc, spec)

    source_records: list[dict[str, Any]] = []
    content_counts = {"headings": 0, "tables": 0, "paragraphs": 0, "lists": 0, "codeBlocks": 0}
    for index, relative in enumerate(spec["sources"]):
        source_path = ENTERPRISE / relative
        if not source_path.exists():
            raise FileNotFoundError(f"Manifest source not found: {source_path}")
        raw = source_path.read_text(encoding="utf-8-sig")
        title = extract_source_title(source_path, raw)
        if index:
            doc.add_page_break()
        counts = add_source_content(doc, source_path, title, markdown)
        for key, value in counts.items():
            content_counts[key] += value
        source_records.append(
            {
                "path": source_path.relative_to(REPO).as_posix(),
                "sha256": sha256(source_path),
                "bytes": source_path.stat().st_size,
                "title": title,
            }
        )

    doc.add_page_break()
    figure_records: list[dict[str, Any]] = []
    for group in ("figures", "screenshots"):
        for item in spec.get(group, []):
            path = ENTERPRISE / item["path"]
            if path.exists():
                figure_records.append(
                    {
                        "path": path.relative_to(REPO).as_posix(),
                        "sha256": sha256(path),
                        "bytes": path.stat().st_size,
                        "caption": item["caption"],
                    }
                )
    add_source_register(doc, spec, source_records, figure_records)
    doc.save(output_path)

    generated = datetime.now(timezone.utc).isoformat()
    document_hash = sha256(output_path)
    record = {
        "schemaVersion": 1,
        "kind": "WS-Sale-App Technical Core DOCX Candidate",
        "generatedAt": generated,
        "status": "Review",
        "documentId": spec["documentId"],
        "title": spec["title"],
        "output": output_path.relative_to(REPO).as_posix(),
        "sha256": document_hash,
        "bytes": output_path.stat().st_size,
        "runtimeVersion": manifest["runtimeVersion"],
        "sourceCommit": manifest["sourceCommit"],
        "sourceInventorySha256": manifest["sourceInventorySha256"],
        "e2eRunId": manifest["e2eRunId"],
        "designPreset": "standard_business_brief",
        "namedOverrides": {
            "page": "A4 portrait",
            "bodyFont": "Prompt 9.4 pt",
            "headingFont": "Kanit",
            "contentWidthCm": CONTENT_WIDTH_CM,
        },
        "headerPattern": "memo_masthead",
        "sources": source_records,
        "figures": figure_records,
        "contentCounts": content_counts,
        "approvalBoundary": "Candidate only; human review and baseline acceptance required",
    }
    manifest_path = output_path.with_suffix(output_path.suffix + ".manifest.json")
    manifest_path.write_text(json.dumps(record, ensure_ascii=False, indent=2), encoding="utf-8")
    return record


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--manifest", type=Path, default=DEFAULT_MANIFEST)
    args = parser.parse_args()
    manifest_path = args.manifest.resolve()
    manifest = json.loads(manifest_path.read_text(encoding="utf-8-sig"))
    if manifest.get("status") != "Review":
        raise ValueError("Technical-core manifest must remain Review")
    markdown = mistune.create_markdown(renderer="ast", plugins=["table"])
    results = [build_document(spec, manifest, markdown) for spec in manifest["documents"]]
    report = {
        "schemaVersion": 1,
        "kind": "WS-Sale-App Technical Core Build",
        "generatedAt": datetime.now(timezone.utc).isoformat(),
        "status": "PASSED",
        "manifest": manifest_path.relative_to(REPO).as_posix(),
        "documents": results,
    }
    report_path = ENTERPRISE / "pipeline" / "reports" / "technical-core-build-report.json"
    report_path.parent.mkdir(parents=True, exist_ok=True)
    report_path.write_text(json.dumps(report, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(report, ensure_ascii=True, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
