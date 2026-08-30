from pathlib import Path
from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\MyWork\WorldFert\winspeed-frontend")
OUT = ROOT / "deliverables" / "cloud-deployment" / "WorldFert_Complete_Operations_Handbook_TH.docx"
OUT.parent.mkdir(parents=True, exist_ok=True)

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
LIGHT_BLUE = "E8EEF5"
PALE_BLUE = "F3F8FC"
GREEN = "DDEFE5"
AMBER = "FFF2CC"
RED = "FCE4E4"
GRAY = "F2F2F2"
TEXT = "24303B"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_text(cell, text, bold=False, color=TEXT, size=9.2, align=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.05
    run = p.add_run(str(text))
    run.bold = bold
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Thai")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_margins(cell)


def add_hyperlink(paragraph, text, url, color=BLUE):
    part = paragraph.part
    rid = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rid)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("หน้า ")
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def add_bullet(doc, text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.left_indent = Inches(0.375 + 0.25 * level)
    p.paragraph_format.first_line_indent = Inches(-0.188)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    if bold_prefix and text.startswith(bold_prefix):
        p.add_run(bold_prefix).bold = True
        p.add_run(text[len(bold_prefix):])
    else:
        p.add_run(text)
    return p


_number_counter = 0
_last_number_element = None


def add_number(doc, text):
    """Add a visible numbered step and restart after any non-number element."""
    global _number_counter, _last_number_element
    previous = None
    for child in reversed(list(doc.element.body)):
        if child.tag != qn("w:sectPr"):
            previous = child
            break
    _number_counter = _number_counter + 1 if previous is _last_number_element else 1
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.375)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    number_run = p.add_run(f"{_number_counter}. ")
    number_run.bold = True
    p.add_run(text)
    _last_number_element = p._p
    return p


def callout(doc, title, body, fill=PALE_BLUE):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_table_header(table.rows[0])
    prevent_row_split(table.rows[0])
    table.columns[0].width = Inches(6.45)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc, headers, rows, widths=None, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, True, "FFFFFF", font_size, WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_shading(table.rows[0].cells[idx], BLUE)
        if widths:
            table.rows[0].cells[idx].width = Inches(widths[idx])
    set_repeat_table_header(table.rows[0])
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        prevent_row_split(table.rows[-1])
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value, False, TEXT, font_size)
            if widths:
                cells[idx].width = Inches(widths[idx])
            if ridx % 2 == 1:
                set_cell_shading(cells[idx], "F8FAFC")
    return table


def heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True
    return p


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.85)
section.bottom_margin = Inches(0.75)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Thai")
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(TEXT)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.22

for style_name in ["List Bullet", "List Bullet 2", "List Number"]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Thai")
    st.font.size = Pt(10.5)

for style_name, size, color, before, after in [
    ("Title", 28, DARK_BLUE, 0, 12),
    ("Heading 1", 16, BLUE, 18, 10),
    ("Heading 2", 13, BLUE, 14, 7),
    ("Heading 3", 12, DARK_BLUE, 10, 5),
]:
    st = styles[style_name]
    st.font.name = "Calibri"
    st._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Thai")
    st.font.size = Pt(size)
    st.font.color.rgb = RGBColor.from_string(color)
    st.font.bold = style_name != "Title" or True
    st.paragraph_format.space_before = Pt(before)
    st.paragraph_format.space_after = Pt(after)

header = section.header
hp = header.paragraphs[0]
hp.text = "WORLDFERT  /  COMPLETE CLOUD VPS OPERATIONS HANDBOOK"
hp.style = styles["Normal"]
hp.runs[0].font.size = Pt(8)
hp.runs[0].font.color.rgb = RGBColor.from_string(BLUE)
hp.runs[0].font.bold = True
footer = section.footer
fp = footer.paragraphs[0]
fp.text = "เอกสารภายใน — WorldFert  |  เวอร์ชัน 2.2  |  "
fp.runs[0].font.size = Pt(8)
fp.runs[0].font.color.rgb = RGBColor.from_string("667788")
add_page_number(fp)

# Cover
doc.add_paragraph("WORLDFERT", style="Subtitle").alignment = WD_ALIGN_PARAGRAPH.LEFT
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(38)
p.paragraph_format.space_after = Pt(12)
r = p.add_run("คู่มือปฏิบัติการ Cloud VPS\nฉบับสมบูรณ์")
r.bold = True
r.font.size = Pt(30)
r.font.color.rgb = RGBColor.from_string(DARK_BLUE)
r.font.name = "Calibri"
r._element.rPr.rFonts.set(qn("w:eastAsia"), "Noto Sans Thai")

p = doc.add_paragraph("สร้าง VPS · ตั้งค่า · Deploy · Restore · เข้าใช้งาน · ทดสอบ · Backup · ดูแลระบบ")
p.paragraph_format.space_after = Pt(26)
p.runs[0].font.size = Pt(15)
p.runs[0].font.color.rgb = RGBColor.from_string(BLUE)

band = doc.add_table(rows=1, cols=3)
band.alignment = WD_TABLE_ALIGNMENT.CENTER
set_repeat_table_header(band.rows[0])
for i, (title, value) in enumerate([
    ("ระบบปัจจุบัน", "Hostinger KVM 2"),
    ("บริการ", "6 Production + 2 Test containers"),
    ("Backup", "อาทิตย์ 02:00 น."),
]):
    cell = band.cell(0, i)
    set_cell_shading(cell, LIGHT_BLUE)
    set_cell_text(cell, f"{title}\n{value}", i == 0, DARK_BLUE, 10.5, WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph("\n")
callout(
    doc,
    "คู่มือหลักฉบับเดียว (Single Source of Truth)",
    "ใช้เอกสารนี้เป็นลำดับปฏิบัติงานตั้งแต่สร้าง VPS จนถึง Deploy, Restore, Login, UAT และงานประจำ. Production-like Pilot ทำงานบน Hostinger KVM 2 (Malaysia), ผูก thirayu.online และใช้ Docker Compose + Portainer โดยไม่ใช้ Coolify หรือ SSH Tunnel. Test Stack แยก Backend/Frontend และฐานข้อมูลพร้อมใช้งานผ่านสคริปต์แบบยืนยันก่อนดำเนินการ.",
    GREEN,
)
p = doc.add_paragraph("ปรับปรุง ณ วันที่ 28 สิงหาคม 2026\nเวอร์ชัน 2.2 · thirayu.online · Test Backend/Frontend · Full Test Stack · Production-to-Test Clone")
p.paragraph_format.space_before = Pt(52)
p.runs[0].font.size = Pt(10)
p.runs[0].font.color.rgb = RGBColor.from_string("667788")

page_break(doc)

heading(doc, "สารบัญการใช้งาน", 1)
contents = [
    ("0", "เริ่มต้นที่นี่: ค่าระบบจริง, บัญชีผู้ใช้, Password key และลำดับงานทั้งหมด"),
    ("1", "บทสรุปสำหรับผู้บริหารและมติที่ต้องอนุมัติ"),
    ("2", "สถาปัตยกรรมเป้าหมายและข้อกำหนด"),
    ("3", "การเตรียม Cloud, DNS, Firewall และเครื่อง Windows"),
    ("4", "ขั้นตอน Deploy ระบบจากระยะไกล"),
    ("5", "การ Upload/Download Backup ผ่าน Direct SFTP"),
    ("6", "การ Restore, Clone Production และ Deploy Test Backend/Frontend แบบแยก"),
    ("7", "การ Backup อัตโนมัติทุกคืนวันอาทิตย์"),
    ("8", "แผน Cutover ภายใน 4 ชั่วโมงและ Rollback"),
    ("9", "การเชื่อมต่อฐานข้อมูลจากภายนอก"),
    ("10", "การทดสอบรับมอบ, Production UAT และ Pilot Monitor 7 วัน"),
    ("11", "เปรียบเทียบ Cloud ต่างประเทศ 5 ราย"),
    ("12", "ตัวเลือกผู้ให้บริการในไทย 5 ราย"),
    ("ภาคผนวก", "รายการ BAT/Shell scripts, checklist และแหล่งอ้างอิง"),
]
add_table(doc, ["ส่วน", "หัวข้อ"], contents, [0.9, 5.5], 9.7)
callout(doc, "วิธีใช้เอกสาร", "เริ่มที่ส่วน 0 เพื่อทราบค่าระบบจริงและตำแหน่ง credentials จากนั้นทำส่วน 3–10 ตามลำดับ. ผู้บริหารอ่านส่วน 1, 8, 11 และ 12. เอกสารนี้ไม่ฝังรหัสผ่านจริงซ้ำ; Password vault ที่ระบุในส่วน 0 เป็นแหล่งข้อมูลที่ได้รับการป้องกัน.")

page_break(doc)
heading(doc, "0. เริ่มต้นที่นี่ — คู่มือหลักฉบับเดียว", 1)
callout(
    doc,
    "กติกาสำคัญก่อนเริ่ม",
    "ห้ามส่ง CREDENTIALS.txt, APPLICATION-ADMIN.txt, private SSH/SFTP key หรือค่า secret จาก .env ผ่าน Chat, Email หรือ Ticket. คู่มือนี้แสดง Username, Password key, ตำแหน่งเก็บ และวิธีเรียกดูอย่างตั้งใจ แต่ไม่ทำสำเนารหัสผ่านจริงเพิ่ม.",
    AMBER,
)

heading(doc, "0.1 ภาพรวมระบบจริง ณ วันที่จัดทำ", 2)
add_table(doc, ["รายการ", "ค่าปัจจุบัน"], [
    ("VPS", "Hostinger VPS ID 1935135 · KVM 2 · Malaysia · Ubuntu 24.04 x86-64"),
    ("ทรัพยากร", "2 vCPU · RAM 8 GB · Disk 100 GB · Swap 2 GB"),
    ("Public IP / Hostname", "76.13.190.104 · srv1935135.hstgr.cloud"),
    ("Domain", "thirayu.online และ app/api/mssql/mysql/portainer/test/api-test subdomains"),
    ("Docker services", "Production 6 ตัว; Test on-demand คือ wf-frontend-test และ wf-backend-test"),
    ("Server application path", "/opt/worldfert/app"),
    ("SFTP transfer path", "/srv/wf-transfer"),
    ("Weekly backup", "ทุกวันอาทิตย์ 02:00 Asia/Bangkok · retention จริงบน Pilot 35 วัน"),
    ("Current status", "Production Pilot deploy/restore/health ผ่าน; Test Stack scripts ผ่าน validation แต่ยังไม่สั่ง Clone/Deploy จริง"),
], [2.0, 4.4], 9.0)

heading(doc, "0.2 Master Access & Credential Registry", 2)
add_table(doc, ["ระบบ/บทบาท", "URL หรือ Host:Port", "Username", "Password / Authentication", "แหล่งข้อมูลจริง"], [
    ("Hostinger hPanel", "https://hpanel.hostinger.com", "บัญชี Hostinger ของเจ้าของระบบ", "Password + MFA ของผู้ใช้; โครงการไม่เก็บ", "Password manager ของเจ้าของบัญชี"),
    ("VPS SSH Admin", "76.13.190.104:22", "root", "Private key: worldfert-hostinger-deploy; root password ผู้ใช้กำหนดเองและไม่ได้เก็บโดยโครงการ", r"deploy\cloud-vps\.local-secrets\worldfert-hostinger-deploy"),
    ("Direct SFTP", "76.13.190.104:22", "wfbackup", "SFTP private key แบบ Ed25519; ไม่มี password", r".local-secrets\worldfert-hostinger-sftp"),
    ("WorldFert App Admin", "https://app.thirayu.online", "บัญชี seed ใน APPLICATION-ADMIN.txt", "Password key: DEFAULT_SEED_PASSWORD", r"deploy\cloud-vps\.local-secrets\APPLICATION-ADMIN.txt"),
    ("WorldFert Test App", "https://test.thirayu.online", "บัญชีที่ถูก clone มาจาก Production", "รหัสผู้ใช้ ณ เวลา Clone; JWT แยกจาก Production", "ฐาน dbwins_worldfert9_test และ test-stack.env บน VPS"),
    ("Portainer Admin", "https://portainer.thirayu.online", "admin", "เรียกดูแบบจงใจ: 11-manage-stack.bat portainer-credentials", "ไฟล์ root-only บน VPS"),
    ("MSSQL Read", "mssql.thirayu.online:1433", "wf_reader", "Password key: WF_READER_PASSWORD", r".local-secrets\CREDENTIALS.txt"),
    ("MSSQL Controlled Write", "mssql.thirayu.online:1433", "wf_owner", "Password key: WF_OWNER_PASSWORD", r".local-secrets\CREDENTIALS.txt"),
    ("MSSQL DBA", "mssql.thirayu.online:1433", "sa", "Password key: MSSQL_SA_PASSWORD", r".local-secrets\CREDENTIALS.txt"),
    ("MySQL Application", "mysql.thirayu.online:3306", "wfapp", "Password key: MYSQL_PASSWORD", r".local-secrets\CREDENTIALS.txt"),
    ("MySQL DBA", "mysql.thirayu.online:3306", "root", "Password key: MYSQL_ROOT_PASSWORD", r".local-secrets\CREDENTIALS.txt"),
    ("Database TLS CA", "ใช้กับ SSMS/DBeaver", "—", "ไม่ใช่ password; เป็น CA certificate", r".local-secrets\worldfert-db-ca.crt"),
], [1.15, 1.20, 1.05, 1.60, 1.40], 7.2)

heading(doc, "0.3 วิธีเปิดดู Credential ที่ได้รับอนุญาต", 2)
add_number(doc, r"บนเครื่องผู้ดูแล เปิดโฟลเดอร์ C:\MyWork\WorldFert\winspeed-frontend\deploy\cloud-vps\.local-secrets.")
add_number(doc, "เปิด CREDENTIALS.txt เฉพาะเมื่อต้องใช้รหัสฐานข้อมูล; ปิดไฟล์หลังใช้งานและห้ามจับภาพหน้าจอ.")
add_number(doc, "เปิด APPLICATION-ADMIN.txt เมื่อต้อง Login ระบบครั้งแรก; Key User ต้องเปลี่ยนรหัสผ่านเริ่มต้นก่อน Production.")
add_number(doc, r"เปิด Command Prompt ที่ deploy\cloud-vps\windows แล้วรัน 11-manage-stack.bat portainer-credentials เมื่อต้อง Login Portainer.")
add_number(doc, "Root password เป็นค่าที่เจ้าของ VPS ตั้งเองและไม่ได้ถูกบันทึกโดยระบบ deploy. หากลืมให้ Reset ผ่าน hPanel ตาม Change Approval; การใช้งานประจำให้ใช้ SSH key.")
add_number(doc, "ห้ามรัน 09-generate-hostinger-profile.ps1 ซ้ำกับระบบใช้งานจริง เพราะจะสร้าง secret ชุดใหม่และอาจทำให้ค่า local ไม่ตรง VPS.")

heading(doc, "0.4 ลำดับดำเนินงานตั้งแต่ศูนย์ถึงพร้อมใช้งาน", 2)
add_table(doc, ["ลำดับ", "งาน", "คำสั่ง/ผลลัพธ์หลัก"], [
    ("1", "สร้าง VPS ใหม่ใน My VPS", "KVM 2 · Malaysia · Ubuntu 24.04 · ตั้ง root password เอง"),
    ("2", "เพิ่ม deploy public key", "ทดสอบ root@Public-IP:22 ด้วย private key"),
    ("3", "Apply Hostinger Firewall", "80/443 Public; 22/1433/3306 เฉพาะ IP /32"),
    ("4", "เตรียม DNS", "@, app, api, mssql, mysql, portainer, test, api-test → Public IP; www → thirayu.online"),
    ("5", "เตรียม Windows profile", "00-check-prerequisites.bat; generator รันครั้งเดียว"),
    ("6", "เตรียม Ubuntu", "01-prepare-server.bat → Docker/SFTP/swap/TLS/directories"),
    ("7", "Deploy empty/current stack", "03-remote-deploy.bat → 6 containers"),
    ("8", "Upload backups", "02-upload-backup.bat → .part + SHA-256 → /incoming"),
    ("9", "Restore databases", "04-restore-mssql.bat และ 05-restore-mysql.bat พร้อม confirmation token"),
    ("10", "ผูก Domain/TLS", "11-manage-stack.bat domain-set thirayu.online"),
    ("11", "ทดสอบทุกช่องทาง", "08-health-check.bat + Browser + SSMS + DBeaver + SFTP + Portainer"),
    ("12", "เปิดงานประจำ", "Weekly backup, Monday download, Pilot monitor, Business UAT"),
], [0.65, 2.15, 3.6], 8.4)

heading(doc, "0.5 ลำดับความสำคัญของแหล่งข้อมูล", 2)
add_bullet(doc, "1) Docker Compose, Caddyfile และ scripts ใน repository คือแหล่ง config แบบ versioned.")
add_bullet(doc, "2) /opt/worldfert/app/.env บน VPS คือค่าที่ระบบกำลังใช้จริง; ตรวจด้วย env-show/env-get.")
add_bullet(doc, "3) .local-secrets คือ credential vault ฝั่งผู้ดูแล; ไม่ commit และไม่ใช้เป็นไฟล์แจกจ่าย.")
add_bullet(doc, "4) คู่มือนี้อธิบายลำดับงานและตำแหน่งค่า; Portainer ใช้ดูสถานะ/ปฏิบัติการ ไม่ใช่แหล่ง config หลัก.")

page_break(doc)
heading(doc, "1. บทสรุปสำหรับผู้บริหาร", 1)
doc.add_paragraph(
    "ระบบรวม Frontend, Backend, Microsoft SQL Server (WinSpeed), MySQL (TruckScale) และ Portainer ไว้บน Cloud VPS เครื่องเดียวด้วย Docker Compose โดยไม่ต้องใช้ Coolify. thirayu.online ถูกผูกกับระบบแล้ว; การรับส่ง backup ใช้ SFTP ผ่าน Public IP โดยตรง และฐานข้อมูลเข้าผ่าน Domain/Public IP ตาม IP allowlist."
)
heading(doc, "1.1 มติที่เสนอให้อนุมัติ", 2)
add_bullet(doc, "รับรอง Hostinger KVM 2 (Malaysia), 2 vCPU / RAM 8 GB / Disk 100 GB เป็น Pilot ราคาต่ำที่ติดตั้งและ restore ระบบครบแล้ว; ใช้ผลวัด 7 วันตัดสินใจสเปก Production.")
add_bullet(doc, "สำหรับ Production ที่มีผู้ใช้จริงพร้อมกันหรือข้อมูลเติบโต แนะนำงบ KVM 4 เป็น baseline และไม่ใช้ราคาโปรโมชันเป็นฐานงบ.")
add_bullet(doc, "อนุมัติให้ Public DB/SFTP ใช้ Source-IP Allowlist ที่ Hostinger Firewall; ไม่เปิด 22/1433/3306 ให้ 0.0.0.0/0. UFW ไม่ได้เปิดใช้งานบน Pilot นี้.")
add_bullet(doc, "อนุมัติ Weekly DB backup คืนวันอาทิตย์ 02:00 น., เก็บบน VPS 35 วัน และดาวน์โหลดคู่ไฟล์ล่าสุดไปเครื่องบริษัท/NAS ผ่าน SFTP ทุกสัปดาห์.")
add_bullet(doc, "อนุมัติช่วง Cutover 4 ชั่วโมง หลังผ่าน Gate: DNS TTL 300, backup ล่าสุดพร้อม, key/config ผ่านการทดสอบ, UAT owner พร้อมตัดสินใจ Go/No-Go.")

page_break(doc)
heading(doc, "1.2 ขอบเขตและข้อจำกัด", 2)
add_table(doc, ["รายการ", "สถานะ/ข้อสรุป"], [
    ("MSSQL backup ปัจจุบัน", "ประมาณ 3.45 GB; ยังต่ำกว่าเพดาน Express 10 GB ต่อฐาน"),
    ("MySQL backup ปัจจุบัน", "ประมาณ 466 MB"),
    ("SQL Server License", "เริ่มด้วย Express ได้; ต้องวางแผนอัปเกรดก่อน database data file แตะ 8 GB"),
    ("Availability", "Single VPS มี Single Point of Failure; provider snapshot และ SFTP copy เป็นชั้นเสริม"),
    ("Backup storage", "ไม่ใช้ S3 ตามข้อกำหนด; ไฟล์อยู่บน VPS และดาวน์โหลดออกผ่าน SFTP"),
    ("เป้าหมายเวลา", "4 ชั่วโมงเป็น Cutover window ไม่รวมงานอนุมัติ/จัดซื้อ/DNS ที่ต้องทำก่อน"),
], [1.8, 4.6], 9.5)
callout(doc, "ความเสี่ยงสำคัญ", "Public IP ช่วยให้เชื่อมต่อง่าย แต่หากเปิดฐานข้อมูลให้ 0.0.0.0/0 จะเพิ่มความเสี่ยงอย่างมีนัยสำคัญ. แผนนี้จึงให้ Public Endpoint เหมือนเดิม แต่จำกัดเฉพาะ IP สำนักงาน/ผู้ใช้งานที่ได้รับอนุมัติ.", AMBER)

page_break(doc)
heading(doc, "2. สถาปัตยกรรมเป้าหมาย", 1)
add_table(doc, ["ชั้น", "องค์ประกอบ", "พอร์ต/เส้นทาง"], [
    ("Edge", "Caddy reverse proxy + TLS อัตโนมัติ", "80, 443 → Public"),
    ("Application", "React/Vite frontend + Node.js backend", "Docker network ภายใน"),
    ("Container management", "Portainer CE LTS ผ่าน Caddy", "443 Public; 9000/9443 ภายในเท่านั้น"),
    ("Database", "SQL Server 2022 + MySQL 8.0", "1433, 3306 → Public + IP allowlist"),
    ("File transfer", "OpenSSH internal-sftp, user wfbackup แบบ chroot", "22 → Public + IP allowlist"),
    ("Operations", "BAT จาก Windows เรียก SSH/SCP/SFTP", "Deploy, restore, backup, download, health check"),
    ("Persistence", "Docker named volumes + /srv/wf-transfer", "DB data แยกจาก release source"),
], [1.2, 3.0, 2.2], 9.2)

heading(doc, "2.1 เส้นทางการเชื่อมต่อ", 2)
add_bullet(doc, "ผู้ใช้เว็บ → app.domain → Caddy:443 → frontend container.")
add_bullet(doc, "Frontend → api.domain → Caddy:443 → backend:3000.")
add_bullet(doc, "ผู้ดูแล → portainer.thirayu.online:443 → Caddy → Portainer:9000 ภายใน; ไม่ publish 9000/9443 สู่ Public IP.")
add_bullet(doc, "Backend → mssql:1433 และ mysql:3306 ผ่าน Docker network; ไม่วิ่งออก Public IP.")
add_bullet(doc, "SSMS/DBeaver/MySQL Workbench → Domain/Public IP → Hostinger Firewall → Database TLS. UFW เป็นชั้นเสริมหากเปิดภายหลัง.")
add_bullet(doc, "Windows Operator → SFTP Public IP:22 → /incoming หรือ /outgoing; ไม่ใช่ SSH tunnel และไม่ใช้ S3.")

heading(doc, "2.2 โครงสร้าง SFTP", 2)
add_table(doc, ["Path ที่ผู้ใช้เห็น", "วัตถุประสงค์", "สิทธิ์"], [
    ("/incoming/mssql", "Upload .bak/.bak.gz + .sha256", "เขียนได้; restore ไม่อัตโนมัติ"),
    ("/incoming/mysql", "Upload .sql/.sql.gz + .sha256", "เขียนได้; restore ไม่อัตโนมัติ"),
    ("/outgoing/mssql", "Download .bak.gz ที่ระบบสร้าง", "อ่าน/ดาวน์โหลด"),
    ("/outgoing/mysql", "Download .sql.gz ที่ระบบสร้าง", "อ่าน/ดาวน์โหลด"),
    ("/manifests", "สถานะ backup ล่าสุด + DB CA certificate", "อ่าน/ดาวน์โหลด"),
], [2.0, 3.1, 1.3], 9.3)

heading(doc, "3. เตรียมระบบก่อน Deploy", 1)
heading(doc, "3.0 สร้าง Hostinger VPS สำหรับ Pilot", 2)
for step in [
    "เข้าสู่ hPanel → VPS → My VPS แล้วเลือกสร้าง VPS ใหม่; ห้ามใช้ Shared VPS ของบัญชีอื่น.",
    "เลือก KVM 2 แบบ 1 เดือน, Region Malaysia, Ubuntu 24.04 LTS และกำหนด root password ด้วยตนเอง. ไม่บันทึก root password ลง repository หรือเอกสารนี้.",
    "รอ VPS เป็น Running แล้วบันทึก Public IPv4 และ Hostname. Pilot นี้คือ 76.13.190.104 และ srv1935135.hstgr.cloud.",
    "ไปที่ Settings → SSH keys แล้วเพิ่ม public key worldfert-hostinger-deploy.pub; ทดสอบ SSH แบบ key-only ก่อนปิดหน้า.",
    "สร้าง Hostinger Firewall ชื่อ WorldFert Production: 22/1433/3306 รับเฉพาะ Public IP ผู้ดูแลแบบ /32, 80/443 รับจาก Any และ Default policy เป็น Drop.",
    "Apply firewall ให้ VPS ที่สร้างใหม่และทดสอบ SSH ทันที. Pilot ใช้ allowlist 58.11.84.165/32; หาก ISP เปลี่ยน IP ต้องแก้ rule ก่อนเชื่อมต่อ.",
]:
    add_number(doc, step)

heading(doc, "3.1 สิ่งที่ต้องได้รับจากผู้ให้บริการ Cloud", 2)
for item in [
    "Ubuntu 24.04 LTS x86-64 พร้อมสิทธิ์ root หรือ sudo; ห้ามเลือก ARM เพราะ SQL Server Linux container ต้องใช้ x86-64.",
    "Public IPv4 แบบคงที่; ยืนยันว่าไม่มี CGNAT และเปิด inbound 22/80/443/1433/3306 ได้.",
    "ขั้นต่ำ Pilot ที่ผ่านการทดสอบ: 2 vCPU, RAM 8 GB, Disk 100 GB; baseline Production: 4 vCPU, RAM 16 GB, NVMe 200 GB และทำ load test ก่อน Go-Live.",
    "Provider Firewall ที่กำหนด Source CIDR ได้ และ snapshot/backup ของ VM เป็นชั้นเสริม.",
    "Region Singapore/SEA หรือประเทศไทย; ทดสอบ latency จากโรงงานและสำนักงานจริง.",
]:
    add_bullet(doc, item)

heading(doc, "3.2 DNS Records", 2)
add_table(doc, ["Name", "Type", "Value / หมายเหตุ"], [
    ("@", "A", "76.13.190.104"),
    ("app", "A", "76.13.190.104"),
    ("api", "A", "76.13.190.104"),
    ("mssql", "A", "76.13.190.104"),
    ("mysql", "A", "76.13.190.104"),
    ("portainer", "A", "76.13.190.104"),
    ("test", "A", "76.13.190.104"),
    ("api-test", "A", "76.13.190.104"),
    ("www", "CNAME", "thirayu.online"),
], [1.2, 1.0, 4.2], 9.2)
doc.add_paragraph("ตั้ง TTL 300 ระหว่าง Pilot/Cutover และคง MX/TXT/AAAA หรือ record อื่นที่ใช้งานอยู่. DNS ไม่กำหนด port; SSMS/DBeaver/SFTP ต้องระบุ 1433/3306/22. หลัง DNS resolve ให้รัน 11-manage-stack.bat domain-set thirayu.online เพื่อ deploy, ออก TLS และหมุน DB certificates.")

heading(doc, "3.3 Firewall Matrix", 2)
add_table(doc, ["Port", "Source", "Action", "เหตุผล"], [
    ("22/TCP", "Office/Admin/SFTP public IP /32", "Allow", "SSH deploy + Direct SFTP"),
    ("80/TCP", "0.0.0.0/0, ::/0", "Allow", "ACME และ redirect"),
    ("443/TCP", "0.0.0.0/0, ::/0", "Allow", "Web/API HTTPS"),
    ("1433/TCP", "Approved client public IP /32", "Allow", "MSSQL TLS"),
    ("3306/TCP", "Approved client public IP /32", "Allow", "MySQL TLS"),
    ("อื่นทั้งหมด", "Any", "Deny", "ลดพื้นที่โจมตี"),
], [0.8, 2.2, 0.9, 2.5], 9.1)

heading(doc, "3.4 เตรียมเครื่อง Windows", 2)
add_number(doc, "เปิด Optional Feature: OpenSSH Client หรือยืนยันว่า ssh.exe, scp.exe และ sftp.exe อยู่ใน PATH.")
add_number(doc, "รัน deploy\\cloud-vps\\windows\\00-check-prerequisites.bat.")
add_number(doc, "สร้าง key คู่ deploy และ SFTP ด้วย ssh-keygen -t ed25519; เก็บ private key เฉพาะเครื่องผู้ดูแล.")
add_number(doc, "เก็บ deploy key และ SFTP key ใน deploy\\cloud-vps\\.local-secrets; ห้าม commit โฟลเดอร์นี้.")
add_number(doc, "รัน PowerShell: 09-generate-hostinger-profile.ps1 -ServerIp 76.13.190.104 -ServerHostname srv1935135.hstgr.cloud -AllowedCidr 58.11.84.165/32 เพื่อสร้าง .env, server-config.env, remote-config.bat และ CREDENTIALS.txt.")
add_number(doc, "ตรวจเฉพาะชื่อ field ในไฟล์ config โดยไม่พิมพ์ secret ลง terminal; ห้ามรัน generator ซ้ำบนระบบใช้งานจริงเพราะจะสร้าง secret ชุดใหม่.")

heading(doc, "4. ขั้นตอน Deploy ระบบจากระยะไกล", 1)
heading(doc, "4.1 Prepare VPS ครั้งแรก", 2)
add_number(doc, "ตั้ง Hostinger Firewall ตามตารางในส่วน 3.3 และ Apply ให้ VPS ก่อนเริ่ม.")
add_number(doc, "เปิด Command Prompt ใน deploy\\cloud-vps\\windows.")
add_number(doc, "รัน 00-check-prerequisites.bat แล้วรัน 01-prepare-server.bat. สคริปต์คง Docker CE ที่ Hostinger ติดตั้งไว้, สร้าง swap 2 GB, สร้าง wfbackup แบบ chroot, private CA และ TLS certificates.")
add_number(doc, "ตรวจ SSH ด้วย root + deploy key และ SFTP ด้วย wfbackup + SFTP key จาก IP ที่อนุญาต. Pilot นี้ใช้ Hostinger Firewall เป็น authoritative control และปล่อย UFW เป็น inactive.")

heading(doc, "4.2 Deploy Application", 2)
add_number(doc, "ตรวจ .env ให้ครบ: Domains, database passwords, JWT/Migrate/Ingress secrets, memory limits.")
add_number(doc, "รัน 03-remote-deploy.bat. สคริปต์สร้าง tar เฉพาะ backend/WSSale-App/deploy, upload ผ่าน SCP, build และ start Docker Compose ใน /opt/worldfert/app.")
add_number(doc, "ระบบตรวจ docker compose config, health ของ wf-mssql/wf-mysql/wf-backend/wf-portainer และติดตั้ง cron คืนวันอาทิตย์. โดยปกติจะคง .env บน VPS; ใช้ --sync-env เฉพาะเมื่อจงใจแทนที่ด้วย local .env.")
add_number(doc, "รัน 08-health-check.bat และเก็บ output ใน Change Record.")

callout(doc, "การ Deploy ซ้ำ", "03-remote-deploy.bat ใช้โฟลเดอร์และ Compose project เดิม จึงคง named volumes ของ MSSQL/MySQL/Caddy ไว้. อย่างไรก็ตาม ให้สร้าง manual backup ก่อน release ที่เปลี่ยน schema ทุกครั้ง.")

page_break(doc)
heading(doc, "4.3 คำสั่งควบคุมหลัก", 2)
add_table(doc, ["BAT", "หน้าที่", "ผลกระทบข้อมูล"], [
    ("00-check-prerequisites.bat", "ตรวจ Windows tools", "ไม่มี"),
    ("01-prepare-server.bat", "เตรียม Ubuntu/SFTP/TLS/swap", "เปลี่ยน server config; ใช้บน VPS ใหม่"),
    ("02-upload-backup.bat", "Upload + SHA-256", "ไม่ restore"),
    ("03-remote-deploy.bat", "Build/Deploy containers", "ไม่ลบ DB volumes"),
    ("04-restore-mssql.bat", "Restore MSSQL + migration + seed", "เขียนทับหลังยืนยัน"),
    ("05-restore-mysql.bat", "Restore MySQL", "เขียนทับหลังยืนยัน"),
    ("06-run-backup-now.bat", "สร้าง backup ทันที", "เพิ่มไฟล์ outgoing"),
    ("07-download-latest-backups.bat", "Download คู่ล่าสุด + verify", "ไม่มีต่อ server"),
    ("08-health-check.bat", "ตรวจ containers/API/web/backup age", "ไม่มี"),
    ("09-generate-hostinger-profile.ps1", "สร้าง config/credentials สำหรับ VPS ใหม่", "สร้าง secret ใหม่; รันครั้งเดียวก่อน deploy"),
    ("10-pilot-monitor.bat", "ติดตั้ง/อ่านรายงาน/หยุด Pilot Monitor", "เก็บเฉพาะ health/ทรัพยากร; ไม่อ่านข้อมูลธุรกิจ"),
    ("11-manage-stack.bat", "จัดการ .env/domain/deploy/rebuild/logs", "backup .env และ validate ก่อนใช้"),
    ("12-clone-databases-to-test.bat", "Clone ฐานปัจจุบันเป็น MSSQL/MySQL Test", "อ่าน Production เพื่อสร้าง snapshot; แทนที่เฉพาะฐานชื่อ test/qa/uat/sandbox หลังยืนยัน"),
    ("13-deploy-test-app.bat", "Build/Rebuild Test Backend และ Frontend", "เปลี่ยนเฉพาะ wf-backend-test/wf-frontend-test; ต้องมีฐาน Test อยู่แล้ว"),
    ("14-prepare-full-test-system.bat", "Clone DB ทั้งสองฐานและ Deploy Test Stack", "แทนที่เฉพาะฐาน Test หลัง safety backup; ไม่ restart Production"),
], [2.1, 2.8, 1.5], 8.5)

heading(doc, "4.4 Operations Manager และ Portainer", 2)
add_table(doc, ["คำสั่ง", "ใช้เมื่อ"], [
    (r"11-manage-stack.bat status | health | connections", "ดูสถานะ, health และ host/port/user โดยไม่แสดง secret"),
    (r"11-manage-stack.bat env-show", "ตรวจ .env โดยปิดบัง PASSWORD/SECRET/TOKEN/PRIVATE_KEY"),
    (r"11-manage-stack.bat env-set APP_VERSION 1.0.1", "แก้ค่าที่ไม่เป็น secret; ยังไม่ deploy"),
    (r"11-manage-stack.bat env-edit", "แก้ทุกค่าแบบ interactive; สำรอง .env และ validate อัตโนมัติ"),
    (r"11-manage-stack.bat deploy", "นำ Compose/config/source ปัจจุบันขึ้นทำงาน"),
    (r"11-manage-stack.bat rebuild backend", "build image ใหม่เมื่อ source/Dockerfile เปลี่ยน"),
    (r"11-manage-stack.bat restart frontend", "restart process โดยไม่ build image"),
    (r"11-manage-stack.bat logs backend 200", "อ่าน 200 บรรทัดล่าสุดเพื่อวิเคราะห์ปัญหา"),
    (r"11-manage-stack.bat env-rollback latest", "ย้อน .env รุ่นล่าสุด แล้วรัน deploy เมื่อยืนยันค่า"),
], [3.5, 2.9], 8.4)
callout(doc, "Portainer", "เข้า https://portainer.thirayu.online ด้วย user admin; ดู password แบบจงใจด้วย 11-manage-stack.bat portainer-credentials. ไม่เปิด 9000/9443, ไม่ให้ผู้ใช้ทั่วไปเป็น admin และให้ Compose/Caddy ใน repository เป็นแหล่ง config หลักเพื่อไม่ให้ UI change สูญหายใน deploy ครั้งถัดไป.", AMBER)

heading(doc, "4.5 ตรวจและแก้ไข .env บน VPS อย่างปลอดภัย", 2)
doc.add_paragraph("Runtime environment ที่ระบบใช้จริงอยู่ที่ /opt/worldfert/app/.env. ไฟล์นี้มี secret และไม่ควรเปิดผ่าน Portainer editor หรือคัดลอกลงเอกสารทั่วไป. Operations Manager จะสำรองไฟล์, ปิดบังค่า secret และตรวจ docker compose config ก่อนนำค่าใหม่ไปใช้.")
add_table(doc, ["งาน", "คำสั่งจาก deploy\\cloud-vps\\windows", "หมายเหตุ"], [
    ("ดูค่าทั้งหมดแบบปิดบัง secret", "11-manage-stack.bat env-show", "แสดง PASSWORD/SECRET/TOKEN/PRIVATE_KEY เป็น redacted"),
    ("ดูค่าเดียว", "11-manage-stack.bat env-get APP_DOMAIN", "ใช้ได้ทุก key; secret จะถูกปิดบัง"),
    ("แก้ค่าที่ไม่ใช่ secret", "11-manage-stack.bat env-set APP_VERSION 1.0.1", "รองรับ allowlist เท่านั้นและยังไม่ deploy"),
    ("แก้ค่าแบบ interactive", "11-manage-stack.bat env-edit", "สร้าง backup, validate และคืนค่าเดิมอัตโนมัติเมื่อไม่ผ่าน"),
    ("ดูรุ่นสำรอง", "11-manage-stack.bat env-backups", "ตรวจ timestamp ก่อน rollback"),
    ("ย้อนค่าล่าสุด", "11-manage-stack.bat env-rollback latest", "ตรวจ env-show แล้ว deploy ใหม่"),
], [1.55, 3.1, 1.75], 8.1)
add_bullet(doc, "Domains: BASE_DOMAIN, APP_DOMAIN, API_DOMAIN, MSSQL_DOMAIN, MYSQL_DOMAIN และ PORTAINER_DOMAIN.")
add_bullet(doc, "Database: ชื่อฐาน, usernames และ password keys ใน Credential Registry; Backend เชื่อมฐานผ่าน Docker network ไม่ผ่าน Public IP.")
add_bullet(doc, "Application secrets: JWT, migration, ingress และ seed password; ต้องเปลี่ยนด้วยขั้นตอนที่อนุมัติและเก็บใน vault.")
add_bullet(doc, "Resource/ports: memory limits, host DB ports และ Docker service versions; เปลี่ยนแล้วต้อง deploy/rebuild ตามประเภท.")
add_bullet(doc, "Backup: CRON, timezone และ retention. ค่าจริงของ Pilot คือ 35 วัน; .env.example อาจแสดง 70 จึงต้องยืนยันด้วย env-get BACKUP_RETAIN_DAYS.")

heading(doc, "4.6 เลือก Deploy, Rebuild, Restart หรือ Rollback", 2)
add_table(doc, ["สถานการณ์", "คำสั่ง", "สิ่งที่ต้องทำก่อน/หลัง"], [
    ("Compose/Caddy/.env เปลี่ยน", "11-manage-stack.bat deploy", "backup-now เมื่อมี schema change; health หลัง deploy"),
    ("Source หรือ Dockerfile เปลี่ยน", "11-manage-stack.bat rebuild backend|frontend", "ตรวจ git version; health และ logs"),
    ("Process ค้างแต่ image/config ไม่เปลี่ยน", "11-manage-stack.bat restart SERVICE", "ดู logs ก่อนและหลัง restart"),
    ("ต้องวิเคราะห์ error", "11-manage-stack.bat logs SERVICE 200", "เก็บ timestamp/correlation; ห้ามนำ secret ลง ticket"),
    ("แก้ .env ผิด", "11-manage-stack.bat env-rollback latest", "ตรวจ env-show → deploy → health"),
    ("Domain/DB certificate เปลี่ยน", "domain-set หรือ rotate-db-certs", "รอ DNS; แจก CA ใหม่ให้ client ที่ได้รับอนุมัติ"),
], [2.05, 2.35, 2.0], 8.2)

heading(doc, "5. Upload/Download Backup ผ่าน Direct SFTP", 1)
heading(doc, "5.1 Upload", 2)
doc.add_paragraph("รูปแบบคำสั่ง:")
add_table(doc, ["ฐาน", "คำสั่ง"], [
    ("MSSQL", r"02-upload-backup.bat mssql C:\Backup\WinSpeed.bak"),
    ("MSSQL gzip", r"02-upload-backup.bat mssql C:\Backup\WinSpeed.bak.gz"),
    ("MySQL", r"02-upload-backup.bat mysql C:\Backup\TruckScale.sql"),
    ("MySQL gzip", r"02-upload-backup.bat mysql C:\Backup\TruckScale.sql.gz"),
], [1.2, 5.2], 9.2)
add_bullet(doc, "BAT คำนวณ SHA-256 ใน Windows และส่งไฟล์ข้อมูลกับ manifest.")
add_bullet(doc, "ไฟล์ขึ้นต้นเป็น .part; เปลี่ยนชื่อเป็นชื่อจริงเมื่อ transfer สำเร็จเท่านั้น.")
add_bullet(doc, "Restore script ปฏิเสธไฟล์ที่ไม่มี manifest หรือ checksum ไม่ตรง.")
add_bullet(doc, "Upload ไม่เรียก restore อัตโนมัติ เพื่อป้องกันการเขียนทับฐานโดยไม่ตั้งใจ.")

heading(doc, "5.2 Download", 2)
add_number(doc, "รัน 07-download-latest-backups.bat.")
add_number(doc, "สคริปต์หาไฟล์ล่าสุดของ MSSQL และ MySQL ผ่าน SSH read-only command.")
add_number(doc, "ดาวน์โหลดไฟล์และ .sha256 ผ่าน SFTP ไป DOWNLOAD_DIR.")
add_number(doc, "PowerShell ตรวจ SHA-256 ฝั่ง Windows; ถ้า mismatch ให้ถือว่า backup ใช้งานไม่ได้.")

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(4)
r = p.add_run("5.3 Client แบบ GUI — ")
r.bold = True
r.font.color.rgb = RGBColor.from_string(BLUE)
p.add_run("WinSCP/FileZilla ใช้ SFTP, Public IP, Port 22, User=wfbackup และ SFTP private key. บัญชีนี้เห็นเฉพาะ chroot folders และไม่มี shell/port forwarding.")

heading(doc, "6. Restore ฐานข้อมูล", 1)
heading(doc, "6.1 MSSQL / WinSpeed", 2)
add_number(doc, "Upload .bak หรือ .bak.gz และ checksum ด้วย 02-upload-backup.bat.")
add_number(doc, "รัน 04-restore-mssql.bat ชื่อไฟล์ และพิมพ์ RESTORE-MSSQL.")
add_number(doc, "ระบบตรวจ checksum, free disk ≥12 GB, RESTORE FILELISTONLY, logical names และเพดาน Express 10 GB.")
add_number(doc, "ระบบรัน RESTORE VERIFYONLY ก่อน. ถ้า backup เก่าไม่มี native checksum จะ restore โดยไม่บังคับ WITH CHECKSUM แต่ยังคง SHA-256 manifest และ VERIFYONLY เป็น gate แล้วจึง SINGLE_USER → RESTORE WITH REPLACE → MULTI_USER.")
add_number(doc, "ระบบสร้าง wf_reader/wf_owner ใหม่, รัน migrations, grant schema wf, seed admin, DBCC UPDATEUSAGE และ DBCC CHECKDB.")

heading(doc, "6.2 MySQL / TruckScale", 2)
add_number(doc, "Upload .sql หรือ .sql.gz และ checksum.")
add_number(doc, "รัน 05-restore-mysql.bat ชื่อไฟล์ และพิมพ์ RESTORE-MYSQL.")
add_number(doc, "ถ้ามี tables เดิม ระบบสร้าง pre-restore .sql.gz ก่อน.")
add_number(doc, "ระบบ drop/create database, import UTF-8, คืนสิทธิ์ wfapp และตรวจจำนวน tables/tblscale.")

heading(doc, "6.3 Clone Production เป็นฐาน Test สำหรับจำลอง Bug", 2)
callout(doc, "หลักความปลอดภัย", "คำสั่ง Clone ไม่เปลี่ยนชื่อฐานใน .env, ไม่ restart Backend และไม่ drop/replace ฐาน Production. สคริปต์ยอมรับฐานปลายทางที่มีคำว่า test, qa, uat หรือ sandbox เท่านั้น และสร้าง safety backup ของฐานทดสอบเดิมก่อนแทนที่.", GREEN)
add_number(doc, r"เปิด Command Prompt ที่ C:\MyWork\WorldFert\winspeed-frontend\deploy\cloud-vps\windows.")
add_number(doc, "รัน 12-clone-databases-to-test.bat โดยไม่ใส่ argument เพื่อ clone ทั้งคู่เป็น dbwins_worldfert9_test และ db_truckscale_test.")
add_number(doc, "สคริปต์รัน read-only preflight ก่อน เพื่อตรวจ container health, ชื่อฐาน, ขนาด และ free disk. ตรวจข้อความให้ถูกต้องแล้วพิมพ์ CLONE-TO-TEST.")
add_number(doc, "MSSQL สร้าง online BACKUP + RESTORE VERIFYONLY, restore ด้วย data/log path ชุดใหม่, map wf_reader/wf_owner และตรวจ DBCC CHECKDB PHYSICAL_ONLY กับจำนวน tables.")
add_number(doc, "MySQL สร้าง consistent mysqldump, import ไปฐาน Test, ปิด scheduled events ที่คัดลอกมา, grant wfapp และตรวจ table count/mysqlcheck.")
add_number(doc, "ตรวจ /srv/wf-transfer/manifests/last-test-clone-status.txt ต้องเป็น status=OK. Safety backup ของฐาน Test เดิมอยู่ใน /outgoing และดาวน์โหลดผ่าน SFTP ได้.")
add_table(doc, ["กรณีใช้งาน", "คำสั่ง"], [
    ("Clone ทั้ง MSSQL และ MySQL ด้วยชื่อมาตรฐาน", "12-clone-databases-to-test.bat"),
    ("Clone ทั้งคู่ด้วยชื่อระบุเอง", "12-clone-databases-to-test.bat all dbwins_bug_123_test db_truckscale_bug_123_test"),
    ("Clone เฉพาะ MSSQL", "12-clone-databases-to-test.bat mssql dbwins_bug_123_test db_truckscale_test"),
    ("Clone เฉพาะ MySQL", "12-clone-databases-to-test.bat mysql dbwins_worldfert9_test db_truckscale_bug_123_test"),
], [2.25, 4.15], 8.4)
add_table(doc, ["Client", "Host:Port", "Database Test", "บัญชีเดิม"], [
    ("SSMS/DBeaver MSSQL", "mssql.thirayu.online:1433", "dbwins_worldfert9_test", "wf_reader / wf_owner / sa"),
    ("DBeaver MySQL", "mysql.thirayu.online:3306", "db_truckscale_test", "wfapp / root"),
], [1.55, 1.85, 1.65, 1.35], 8.3)
callout(doc, "การทดสอบผ่าน Application", "ห้ามเปลี่ยน Backend สาธารณะให้ชี้ฐาน Test เพราะผู้ใช้ Production จะถูกส่งไปยังข้อมูลทดสอบ. ใช้ Test Backend/Frontend แยก URL ตามส่วน 6.4 ซึ่งกำหนด JWT แยก, ปิด LINE/Webhook และปิด Background Workers โดยค่าเริ่มต้น.", AMBER)

heading(doc, "6.4 Test Backend/Frontend และ Full Test Stack", 2)
callout(doc, "สถาปัตยกรรม Test ที่แยกชัดเจน", "wf-backend-test และ wf-frontend-test เป็นคนละ container/image/volume จาก Production แต่เชื่อม wf-mssql และ wf-mysql ผ่าน Docker network เดิม โดยระบุเฉพาะ dbwins_worldfert9_test และ db_truckscale_test. Production containers และชื่อฐาน Production ไม่ถูก restart หรือแก้ไข.", GREEN)
add_table(doc, ["ส่วน", "ชื่อ/URL", "การแยกจาก Production"], [
    ("Test Frontend", "https://test.thirayu.online · wf-frontend-test", "Build ด้วย api-test URL และแสดงป้าย TEST SYSTEM"),
    ("Test API", "https://api-test.thirayu.online/api · wf-backend-test", "JWT แยก; CORS เฉพาะ test domain"),
    ("Test MSSQL", "dbwins_worldfert9_test", "ชื่อฐานแยกใน wf-mssql"),
    ("Test MySQL", "db_truckscale_test", "ชื่อฐานแยกใน wf-mysql; copied events ถูกปิด"),
], [1.35, 2.45, 2.60], 8.3)
add_number(doc, "สร้าง DNS A records test และ api-test ให้ชี้ 76.13.190.104; Caddy จะออก HTTPS certificate เมื่อ DNS พร้อม.")
add_number(doc, "ถ้ามีฐาน Test อยู่แล้วและต้องการ Deploy/Rebuild เฉพาะแอป ให้รัน 13-deploy-test-app.bat แล้วพิมพ์ DEPLOY-TEST-APP.")
add_number(doc, "ถ้าต้องการข้อมูล Production ล่าสุดพร้อมระบบ Test ทั้งชุด ให้รัน 14-prepare-full-test-system.bat แล้วพิมพ์ REBUILD-FULL-TEST. ระบบจะ preflight, หยุดเฉพาะ Test app, safety-backup ฐาน Test เดิม, Clone ฐานทั้งสอง และ Deploy Test app.")
add_number(doc, "เปิด https://test.thirayu.online ตรวจป้าย TEST SYSTEM, Login ด้วยบัญชี ณ เวลา Clone และเปิด https://api-test.thirayu.online/api/health เพื่อตรวจ sqlserver=mysql=up.")
add_number(doc, "ตรวจ Portainer ต้องเห็น wf-backend-test และ wf-frontend-test แยกจาก wf-backend/wf-frontend. Test backend ใช้ NODE_ENV=test, JWT แยก, LINE/Webhook ว่าง และ DISABLE_BACKGROUND_WORKERS=true.")
add_table(doc, ["งาน", "คำสั่ง"], [
    ("Deploy เฉพาะ Test app", "13-deploy-test-app.bat"),
    ("Full refresh ด้วยชื่อมาตรฐาน", "14-prepare-full-test-system.bat"),
    ("Full refresh ตาม Bug", "14-prepare-full-test-system.bat dbwins_bug_123_test db_truckscale_bug_123_test"),
    ("ดูสถานะผ่าน SSH", "sudo APP_DIR=/opt/worldfert/app .../server/deploy-test-app.sh status"),
    ("ดู Backend Test logs", "sudo APP_DIR=/opt/worldfert/app .../server/deploy-test-app.sh logs backend-test 200"),
    ("หยุดเฉพาะ Test app", "sudo APP_DIR=/opt/worldfert/app .../server/deploy-test-app.sh stop"),
], [2.1, 4.3], 8.0)
callout(doc, "ข้อควรระวังด้านข้อมูลส่วนบุคคล", "ฐาน Test เป็นสำเนาข้อมูล Production จึงต้องจำกัดผู้เข้าถึงเทียบเท่า Production, ห้ามส่งไฟล์ออกนอกองค์กร และควร Mask ข้อมูลส่วนบุคคลเมื่อเปิดให้ Developer/บุคคลที่ไม่จำเป็นต้องเห็นข้อมูลจริง. การหยุด Test app จะคงฐานและ volumes ไว้; ให้กำหนด retention/ลบตาม Change Approval.", AMBER)

heading(doc, "6.5 เกณฑ์หยุดและไม่ฝืน Restore/Clone/Test Deploy", 2)
add_table(doc, ["เหตุการณ์", "การตัดสินใจ"], [
    ("Checksum ไม่ตรง/ไม่มี manifest", "หยุด; upload ใหม่ ห้าม bypass"),
    ("MSSQL data ≥10 GB บน Express", "หยุด; เปลี่ยน Edition/License ก่อน"),
    ("Free disk <12 GB", "หยุด; เพิ่ม disk/ลบไฟล์ที่อนุมัติแล้ว"),
    ("DBCC CHECKDB มี error", "No-Go; เก็บหลักฐานและกลับระบบเดิม"),
    ("tblscale ไม่มี/rows ผิดปกติ", "No-Go; ตรวจว่าใช้ dump ถูกฐาน"),
    ("Migration/seed ล้มเหลว", "No-Go; ใช้ pre-restore backup หรือ rollback VM"),
    ("Clone target ไม่มี test/qa/uat/sandbox", "หยุด; เปลี่ยนชื่อฐาน Test ห้าม bypass validation"),
    ("พื้นที่ไม่ผ่าน Clone preflight", "หยุด; เพิ่ม disk หรือจัดการไฟล์ที่อนุมัติก่อน ห้ามลด safety margin"),
    ("Available RAM ต่ำกว่า 1 GB", "หยุด Test Deploy; ตรวจ capacity/container limits ก่อน"),
    ("Test API DB status ไม่เป็น up ทั้งคู่", "ห้ามส่งให้ผู้ทดสอบ; ตรวจชื่อฐาน, grants, TLS และ Test Backend logs"),
], [2.4, 4.0], 9.3)

heading(doc, "7. Backup รายสัปดาห์", 1)
doc.add_paragraph("ตารางอัตโนมัติ: ทุกวันอาทิตย์ 02:00 น. ตาม Asia/Bangkok; Hostinger Pilot เก็บบน VPS 35 วัน (ประมาณ 5 รุ่น) และให้ดาวน์โหลดไป NAS/เครื่องบริษัทเพื่อเก็บระยะยาว.")
add_table(doc, ["ขั้น", "MSSQL", "MySQL"], [
    ("สร้าง", "BACKUP DATABASE WITH CHECKSUM; fallback เมื่อ Express ไม่รองรับ compression", "mysqldump --single-transaction --quick + routines/triggers/events"),
    ("ตรวจ", "RESTORE VERIFYONLY", "gzip -t และขนาดขั้นต่ำ"),
    ("บีบอัด", ".bak.gz", ".sql.gz"),
    ("ยืนยัน", "SHA-256 manifest", "SHA-256 manifest"),
    ("ปลายทาง", "/outgoing/mssql", "/outgoing/mysql"),
], [1.0, 2.8, 2.8], 9.1)
heading(doc, "7.1 ขั้นตอน Operator ทุกสัปดาห์", 2)
add_bullet(doc, "เช้าวันจันทร์ ตรวจ /manifests/last-backup-status.txt ว่า status=OK และ completed_at ล่าสุด.")
add_bullet(doc, "รัน 07-download-latest-backups.bat เพื่อเก็บสำเนานอก VPS; ห้ามถือว่าไฟล์บน VPS เครื่องเดียวเป็น Disaster Recovery copy.")
add_bullet(doc, "เก็บ local/NAS อย่างน้อย 10 รุ่น และจำกัดสิทธิ์เฉพาะทีม IT.")
add_bullet(doc, "ทำ Restore Drill รายไตรมาสบนเครื่องทดสอบ; บันทึกเวลาจริงและผล DBCC/tblscale.")
heading(doc, "7.2 Backup ก่อนเปลี่ยนระบบ", 2)
doc.add_paragraph("ก่อน deploy release หรือ restore ซ้ำ ให้รัน 06-run-backup-now.bat. Restore scripts สร้าง pre-restore backup อัตโนมัติเมื่อพบฐานเดิม แต่ manual backup ยังเป็นหลักฐานก่อน Change Window ที่ชัดเจนกว่า.")

page_break(doc)
heading(doc, "8. แผน Cutover ภายใน 4 ชั่วโมง", 1)
callout(doc, "เงื่อนไขสำเร็จ", "4 ชั่วโมงทำได้เมื่อ Cloud/DNS/Firewall/key/.env ถูกเตรียมล่วงหน้า, DNS TTL=300 และอัปโหลด 3.9 GB ได้ภายในประมาณ 30 นาที. หาก uplink ต่ำกว่า ~20 Mbps ให้ pre-stage full backup แล้วส่งเฉพาะ final delta/backup ล่าสุดใน window.", AMBER)
add_table(doc, ["เวลา", "กิจกรรม", "Owner", "Gate/ผลลัพธ์"], [
    ("T-24h ถึง T-2h", "Provision VPS, DNS TTL 300, provider firewall, prepare/deploy empty stack, pre-stage backup", "Infra", "Health check ผ่าน; restore ยังไม่ทำ"),
    ("00:00–00:15", "ประกาศเริ่ม, freeze การเขียน, บันทึกยอด/เวลา, สร้าง final backups", "Business + DBA", "ยืนยัน Write Freeze"),
    ("00:15–00:45", "Upload final .bak/.sql + SHA-256 ผ่าน SFTP", "DBA", "Checksum ผ่าน"),
    ("00:45–01:30", "Restore MSSQL + DBCC + migrations/seed", "DBA/App", "MSSQL ผ่าน"),
    ("01:30–01:55", "Restore MySQL + grant + row checks", "DBA", "MySQL ผ่าน"),
    ("01:55–02:25", "Deploy/rebuild final release + health checks", "App/Infra", "Containers/API healthy"),
    ("02:25–03:10", "UAT: Login, WinSpeed read, quotation/order, TruckScale, export/report", "Key Users", "Sign-off หรือ Rollback"),
    ("03:10–03:30", "เปิด DNS/ผู้ใช้จริง, ทดสอบ Public DB clients", "Infra", "Go-Live"),
    ("03:30–04:00", "Monitor logs/CPU/RAM/disk, ดาวน์โหลด post-cutover backup", "War room", "Close หรือ rollback"),
], [1.05, 2.75, 1.15, 1.65], 8.4)

page_break(doc)
heading(doc, "8.1 Go/No-Go ณ นาทีที่ 190", 2)
add_bullet(doc, "Go เมื่อ DBCC ผ่าน, tblscale ถูกต้อง, API healthy, UAT critical flows ผ่าน และไม่มี error rate/latency ผิดปกติ.")
add_bullet(doc, "No-Go เมื่อข้อมูลไม่ครบ, migration fail, DB connection ไม่เสถียร หรือเหลือเวลาน้อยกว่า 30 นาทีโดยยังไม่ผ่าน critical UAT.")

heading(doc, "8.2 Rollback", 2)
add_number(doc, "หยุดรับผู้ใช้บน Cloud; เก็บ logs และเวลาที่ตัดสินใจ.")
add_number(doc, "คืน DNS/entry point ไประบบเดิม. TTL 300 ช่วยลดเวลารอ.")
add_number(doc, "ยกเลิก Write Freeze บนระบบเดิมเมื่อยืนยันว่าไม่มี write ใหม่บน Cloud หรือ reconciliation เสร็จ.")
add_number(doc, "เก็บ Cloud volumes/backup ไว้ห้ามลบ เพื่อทำ Root Cause Analysis.")
add_number(doc, "นัด Change Window ใหม่หลังแก้สาเหตุและผ่าน Restore Drill.")

heading(doc, "9. การเชื่อมต่อฐานข้อมูลจากภายนอก", 1)
heading(doc, "9.1 ค่าการเชื่อมต่อจริงและตำแหน่งรหัสผ่าน", 2)
add_table(doc, ["ระบบ", "Host:Port / Database", "Username", "Password key"], [
    ("MSSQL อ่านอย่างเดียว", "mssql.thirayu.online:1433\ndbwins_worldfert9", "wf_reader", "WF_READER_PASSWORD"),
    ("MSSQL เขียนแบบควบคุม", "mssql.thirayu.online:1433\ndbwins_worldfert9", "wf_owner", "WF_OWNER_PASSWORD"),
    ("MSSQL Admin", "mssql.thirayu.online:1433\ndbwins_worldfert9", "sa", "MSSQL_SA_PASSWORD"),
    ("MySQL App", "mysql.thirayu.online:3306\ndb_truckscale", "wfapp", "MYSQL_PASSWORD"),
    ("MySQL Admin", "mysql.thirayu.online:3306\ndb_truckscale", "root", "MYSQL_ROOT_PASSWORD"),
], [1.35, 2.15, 1.25, 1.65], 8.5)
doc.add_paragraph(r"ค่าจริงของ Password อยู่ในไฟล์ที่ไม่ commit: deploy\cloud-vps\.local-secrets\CREDENTIALS.txt. เอกสารทั่วไปไม่ฝังรหัสผ่านเพื่อป้องกันการส่งต่อโดยไม่ตั้งใจ. ใช้ wf_reader สำหรับตรวจสอบ/รายงาน, wf_owner เฉพาะงานเขียนที่อนุมัติ และ sa/root เฉพาะ DBA.")

heading(doc, "9.2 SSMS — MSSQL", 2)
add_number(doc, "เปิด Connect → Database Engine; Server name = mssql.thirayu.online,1433; Authentication = SQL Server Authentication.")
add_number(doc, "ใส่ login ตามหน้าที่และเลือก Database = dbwins_worldfert9.")
add_number(doc, r"ติดตั้ง deploy\cloud-vps\.local-secrets\worldfert-db-ca.crt ใน Windows Trusted Root Certification Authorities (Local Computer หรือ Current User ตามนโยบายองค์กร).")
add_number(doc, "Connection Security: Encryption = Mandatory และไม่เลือก Trust server certificate; กด Connect แล้วตรวจว่า certificate verification ผ่าน.")
add_number(doc, "Trust server certificate=True ใช้ชั่วคราวเฉพาะ migration window เมื่อ IP allowlist ทำงาน และต้องกลับเป็น False หลังติดตั้ง CA.")

heading(doc, "9.3 DBeaver — MSSQL และ MySQL", 2)
add_bullet(doc, "MSSQL: SQL Server driver; Host=mssql.thirayu.online, Port=1433, Database=dbwins_worldfert9, SQL login ตามหน้าที่; เปิด encrypt/certificate validation และ trustServerCertificate=false.")
add_bullet(doc, "MySQL: Host=mysql.thirayu.online, Port=3306, Database=db_truckscale, User=wfapp; SSL mode=VERIFY_CA หรือ VERIFY_IDENTITY และเลือก worldfert-db-ca.crt เป็น CA Certificate.")
add_bullet(doc, "หาก DBeaver/Java ไม่ใช้ Windows certificate store ให้เพิ่ม CA ใน SSL/driver trust configuration แล้ว Test Connection; ห้ามยอมรับ unknown certificate ถาวร.")
add_bullet(doc, "Fallback ชั่วคราว: MySQL REQUIRED หรือ MSSQL trustServerCertificate=true ระหว่าง migration เท่านั้น; ทั้งสองแบบเข้ารหัสแต่ลดการยืนยันตัวตนของ server.")

heading(doc, "9.4 เมื่อ Public IP ผู้ใช้เปลี่ยน", 2)
doc.add_paragraph("อัปเดต Source CIDR ใน Hostinger/Provider Firewall แล้วทดสอบพอร์ตจาก client; หากภายหลังเปิด UFW ให้แก้ CIDR ให้สอดคล้องกันด้วย. ห้ามแก้เป็น 0.0.0.0/0 เพื่อแก้ปัญหาชั่วคราวโดยไม่มี Change Approval. ถ้าจำเป็นจริง ค่า ALLOW_DATABASES_FROM_ANYWHERE=true ต้องถูกเปิดอย่างจงใจและกำหนดเวลาปิดกลับ.")

heading(doc, "9.5 การเข้าใช้งานระบบอื่น", 2)
add_table(doc, ["ระบบ", "วิธีเข้า", "เกณฑ์ทดสอบ"], [
    ("Frontend", "https://app.thirayu.online · ใช้บัญชีใน APPLICATION-ADMIN.txt ครั้งแรก", "Login, Dashboard, ภาษาไทย, timezone และ critical business flow"),
    ("API", "https://api.thirayu.online/api/health", "HTTP 200 และ MSSQL/MySQL = up; endpoint อื่นต้องผ่านสิทธิ์"),
    ("Portainer", "https://portainer.thirayu.online · admin · เรียก password ด้วยคำสั่งที่ระบุ", "เห็น 6 containers; 9000/9443 ต้องไม่เปิด Public"),
    ("SFTP", "76.13.190.104:22 · wfbackup · worldfert-hostinger-sftp", "เห็น /incoming, /outgoing, /manifests; ไม่มี shell/forwarding"),
], [1.15, 3.35, 1.9], 8.2)

heading(doc, "10. Acceptance Test และงานหลัง Go-Live", 1)
heading(doc, "10.1 Technical Acceptance", 2)
checks = [
    "[ ] docker compose ps แสดง wf-caddy/wf-frontend/wf-backend/wf-mssql/wf-mysql/wf-portainer ทำงานครบ",
    "[ ] https://APP_DOMAIN และ https://API_DOMAIN/api/health ผ่านจากภายนอก",
    "[ ] https://portainer.thirayu.online แสดงหน้า Login และ host ไม่ listen 9000/9443",
    "[ ] MSSQL DB ONLINE, collation ถูกต้อง, DBCC UPDATEUSAGE และ DBCC CHECKDB ไม่พบ error",
    "[ ] MySQL tables และ tblscale rows อยู่ในช่วงที่คาด",
    "[ ] SSMS และ MySQL client ต่อด้วย TLS จาก IP ที่อนุญาต และถูกปฏิเสธจาก IP อื่น",
    "[ ] SFTP upload/download ผ่านด้วย wfbackup แต่ shell/port forwarding ใช้ไม่ได้",
    "[ ] 06-run-backup-now และ 07-download-latest-backups ผ่าน checksum",
    "[ ] Disk free หลัง backup ≥25%; memory ไม่มี OOM; container ไม่มี restart loop",
]
for c in checks:
    doc.add_paragraph(c)

heading(doc, "10.2 Business UAT", 2)
for c in [
    "[ ] Login admin และผู้ใช้จริง",
    "[ ] อ่านข้อมูล WinSpeed สำคัญและค้นหาเอกสารย้อนหลัง",
    "[ ] สร้าง/อนุมัติ/พิมพ์เอกสารธุรกิจที่อยู่ใน critical flow",
    "[ ] TruckScale แสดงรายการ/สถานะ/น้ำหนักถูกต้อง และไม่เกิด duplicate write",
    "[ ] Export/report ภาษาไทยและเวลา Asia/Bangkok ถูกต้อง",
    "[ ] LINE Login (ถ้าใช้งาน) callback และ redirect ถูก domain",
]:
    doc.add_paragraph(c)

heading(doc, "10.3 ผล Hostinger Pilot ที่ยืนยันแล้ว", 2)
add_table(doc, ["รายการ", "ผลจริง ณ 28 ส.ค. 2026"], [
    ("VPS", "ID 1935135 · KVM 2 · Malaysia · Ubuntu 24.04 · 2 vCPU · RAM 8 GB · Disk 100 GB"),
    ("ราคา", "$14.97 งวดแรก 1 เดือนรวมภาษี · Renewal $24.49/เดือน · หมดอายุ 27 ก.ย. 2026 · Auto-renew เปิด"),
    ("Web/API/Domain", "thirayu.online, app.thirayu.online และ api.thirayu.online/api/health ผ่าน HTTPS; root redirect เข้า app และ API รายงาน MSSQL/MySQL = up"),
    ("Portainer", "portainer.thirayu.online ผ่าน HTTPS; admin ถูกสร้างด้วย root-only password file; ไม่มี public listener 9000/9443"),
    ("MSSQL", "dbwins_worldfert9 restore สำเร็จ · 734 tables · migrations 073–101 สำเร็จ · DBCC CHECKDB ผ่าน"),
    ("MySQL", "db_truckscale restore สำเร็จ · 21 tables · tblscale 403,908 rows"),
    ("Public DB", "mssql.thirayu.online:1433 และ mysql.thirayu.online:3306 มี DNS/TLS SAN ใหม่; legacy hostname/IP ยังคงเป็น certificate aliases และ client TLS + CA verification ผ่าน"),
    ("SFTP/Backup", "wfbackup key-only ผ่าน · manual backup/download/ตรวจ SHA-256 ผ่าน · cron อาทิตย์ 02:00 · retention 35 วัน"),
    ("Production UAT", "11/11 API ขั้นผ่าน; หน้า Login และ Dashboard จริงผ่านจาก browser พร้อมข้อมูล SO 125,179 รายการ; flow สร้าง/อนุมัติเอกสารยังรอ Key User"),
    ("ทรัพยากรล่าสุด", "หลังเพิ่ม Portainer: RAM available 5.5 GiB · Swap ใช้ 512 KiB · Disk ใช้ 17/96 GiB (18%), เหลือ 80 GiB"),
], [1.55, 4.85], 8.9)
callout(doc, "สถานะรับมอบทางเทคนิค", "Containers ทั้ง 6 ทำงาน, HTTPS/Public DB/SFTP/backup ผ่าน, Portainer พร้อม Login และ Production-compatible API UAT ผ่านครบ. Browser Login/Dashboard ผ่าน; Business UAT ที่สร้างหรืออนุมัติเอกสารจริง และการเปลี่ยนรหัส admin เริ่มต้นยังต้องให้ Key User ดำเนินการก่อน Production.", GREEN)

heading(doc, "10.4 Pilot Monitor 7 วัน", 2)
doc.add_paragraph("เริ่ม 28 ส.ค. 2026 09:16 น. และวางแผนจบ 4 ก.ย. 2026 09:16 น. เก็บตัวอย่างทุก 5 นาทีที่ /var/log/worldfert-pilot/metrics.csv โดยไม่อ่านหรือแก้ข้อมูลธุรกิจ.")
add_bullet(doc, r"เริ่ม/ดูรายงาน/หยุด: windows\10-pilot-monitor.bat install 7 | report | stop")
add_bullet(doc, "แจ้งเตือนเมื่อ container/API/DB มี failed sample, Disk >75%, RAM available <1 GB หรือ backup status ไม่ใช่ OK.")
add_bullet(doc, "รายงานล่าสุดที่บันทึกไว้: 61 samples, max load1 1.74, min RAM available 3,465 MB, max Disk 18%, unhealthy/API/DB failed samples = 0 และ weekly backup status = OK.")
add_bullet(doc, "Codex heartbeat ตรวจวันละครั้ง 7 รอบ; เมื่อครบกำหนดให้สรุปผลก่อน stop cron และคง metrics ไว้เป็นหลักฐาน.")

heading(doc, "10.5 งาน 7/30/90 วัน", 2)
add_table(doc, ["กำหนด", "งาน"], [
    ("ภายใน 7 วัน", "ติดตาม CPU/RAM/disk/latency, จำกัด login ภายนอก, ติดตั้ง DB CA ทุก client"),
    ("ภายใน 30 วัน", "Restore Drill, ทบทวน firewall, rotate temporary credentials, ประเมิน disk growth"),
    ("ภายใน 90 วัน", "Load test, capacity review, ตัดสินใจแยก DB server/HA และทบทวน SQL Edition"),
], [1.2, 5.2], 9.4)

heading(doc, "10.6 ตารางงานประจำของ Operator", 2)
add_table(doc, ["ความถี่", "งานขั้นต่ำ", "หลักฐานที่ต้องเก็บ"], [
    ("ทุกวันทำการ", "status → health; ตรวจ restart count และ error logs เมื่อผิดปกติ", "ผล health พร้อมเวลา"),
    ("เช้าวันจันทร์", "ตรวจ last-backup-status.txt = OK และดาวน์โหลด backup คู่ล่าสุด", "ไฟล์ MSSQL/MySQL + .sha256 บน NAS"),
    ("ก่อนเปลี่ยนระบบ", "backup-now; env-backups; บันทึก release/version และ rollback owner", "Change Record + backup timestamp"),
    ("ทุกเดือน", "ทบทวน firewall /32, users, disk growth, certificate และ credential ownership", "Monthly access/capacity review"),
    ("ทุกไตรมาส", "Restore Drill บนระบบทดสอบ; DBCC/tables/rows/UAT", "Restore evidence และเวลาที่ใช้จริง"),
], [1.15, 3.55, 1.7], 8.3)

heading(doc, "10.7 ขั้นตอนมาตรฐานเมื่อปรับปรุงระบบ", 2)
for step in [
    "เปิด Change Record: ระบุเหตุผล, release, ผู้อนุมัติ, downtime, rollback owner และเวลาที่คาด.",
    "รัน 08-health-check.bat และ 06-run-backup-now.bat; ยืนยัน backup status=OK ก่อนแก้ระบบ.",
    "ตรวจค่า runtime ด้วย env-show/env-get. หากต้องแก้ secret ให้ใช้ env-edit ใน session ที่ปลอดภัยและไม่บันทึก output.",
    "เลือก deploy/rebuild/restart ตามตาราง 4.6. ห้ามลบ Docker volumes หรือใช้ docker compose down -v.",
    "รัน health, ตรวจ logs และทดสอบ Browser/API/DB/SFTP เฉพาะส่วนที่ได้รับผลกระทบ.",
    "ถ้าไม่ผ่าน ให้ rollback .env/release หรือคืน entry point ระบบเดิมตามแผน; เก็บ volumes/logs ไว้ตรวจสาเหตุ.",
    "ปิด Change Record เมื่อ Key User ยืนยันและไม่มี error/restart ผิดปกติในช่วงเฝ้าระวัง.",
]:
    add_number(doc, step)

heading(doc, "10.8 Troubleshooting Decision Table", 2)
add_table(doc, ["อาการ", "ตรวจตามลำดับ", "ข้อห้าม/ทางออก"], [
    ("SSH/SFTP เข้าไม่ได้", "Public IP ผู้ใช้ → Hostinger Firewall /32 → port 22 → username/key → known_hosts", "ห้ามเปิด 22 เป็น Any; reset root password ผ่าน hPanel เมื่อจำเป็น"),
    ("SSMS/DBeaver ต่อ DB ไม่ได้", "DNS → Firewall 1433/3306 → CA/SSL mode → username/password key → DB health", "ห้ามใช้ Trust server certificate ถาวร; ห้ามเปิด DB 0.0.0.0/0"),
    ("Web 502 / API unhealthy", "status → health → logs backend/caddy → DB health → disk/RAM", "restart เฉพาะเมื่อบันทึก logs แล้ว; rollback ถ้าเกิดหลัง release"),
    ("Portainer เข้าไม่ได้", "HTTPS/DNS → wf-portainer status/logs → portainer-restart", "ห้าม publish 9000/9443 เพื่อแก้ชั่วคราว"),
    (".env validate ไม่ผ่าน", "env-backups → env-show → env-rollback latest → deploy", "ห้ามแทนที่ remote .env ด้วย --sync-env โดยไม่ตรวจ diff"),
    ("Backup status ไม่ใช่ OK", "manifest → backup logs → disk free → DB health → backup-now", "ถือว่าไม่มี backup จนสร้างและ verify ใหม่สำเร็จ"),
    ("Disk >75% หรือ RAM <1 GB", "ดู backup retention/log growth/container usage → capacity review", "ห้ามลบ DB volume; archive/delete เฉพาะไฟล์ที่อนุมัติ"),
], [1.35, 3.25, 1.8], 7.8)

page_break(doc)
heading(doc, "10.9 Command Cheat Sheet — ใช้จาก Windows", 2)
doc.add_paragraph(r"เปิด Command Prompt ที่ C:\MyWork\WorldFert\winspeed-frontend\deploy\cloud-vps\windows แล้วใช้คำสั่งต่อไปนี้:")
add_table(doc, ["เป้าหมาย", "คำสั่ง"], [
    ("ตรวจ prerequisites", "00-check-prerequisites.bat"),
    ("เตรียม VPS ใหม่", "01-prepare-server.bat"),
    ("Upload MSSQL", r"02-upload-backup.bat mssql C:\Backup\WinSpeed.bak"),
    ("Upload MySQL", r"02-upload-backup.bat mysql C:\Backup\TruckScale.sql"),
    ("Deploy application", "03-remote-deploy.bat"),
    ("Restore MSSQL", "04-restore-mssql.bat <filename>  แล้วพิมพ์ RESTORE-MSSQL"),
    ("Restore MySQL", "05-restore-mysql.bat <filename>  แล้วพิมพ์ RESTORE-MYSQL"),
    ("Backup ทันที", "06-run-backup-now.bat"),
    ("Download backup ล่าสุด", "07-download-latest-backups.bat"),
    ("Health check", "08-health-check.bat"),
    ("Pilot report", "10-pilot-monitor.bat report"),
    ("สถานะ/การเชื่อมต่อ", "11-manage-stack.bat status | health | connections"),
    ("ดู/แก้ .env", "11-manage-stack.bat env-show | env-get KEY | env-set KEY VALUE | env-edit"),
    ("Deploy/Rebuild/Restart", "11-manage-stack.bat deploy | rebuild SERVICE | restart SERVICE"),
    ("Logs", "11-manage-stack.bat logs SERVICE 200"),
    ("Rollback .env", "11-manage-stack.bat env-backups  แล้ว env-rollback latest"),
    ("Portainer password", "11-manage-stack.bat portainer-credentials"),
    ("ตั้ง Domain", "11-manage-stack.bat domain-set thirayu.online"),
    ("Clone Production → Test", "12-clone-databases-to-test.bat  แล้วพิมพ์ CLONE-TO-TEST"),
    ("Deploy Test Backend/Frontend", "13-deploy-test-app.bat  แล้วพิมพ์ DEPLOY-TEST-APP"),
    ("Refresh Test ทั้งระบบ", "14-prepare-full-test-system.bat  แล้วพิมพ์ REBUILD-FULL-TEST"),
], [2.1, 4.3], 8.2)
callout(doc, "คำสั่งที่ห้ามใช้โดยไม่มีแผนกู้คืน", "ห้าม docker compose down -v, docker volume rm, ลบ /opt/worldfert/app/.env, ลบ /srv/wf-transfer หรือสร้าง profile ใหม่ทับของเดิม. การกระทำเหล่านี้อาจลบข้อมูลหรือทำให้ credentials ไม่ตรงกัน.", RED)

# Landscape comparison
land = doc.add_section(WD_SECTION.NEW_PAGE)
land.orientation = WD_ORIENT.LANDSCAPE
land.page_width = Inches(11)
land.page_height = Inches(8.5)
land.top_margin = Inches(0.65)
land.bottom_margin = Inches(0.65)
land.left_margin = Inches(0.65)
land.right_margin = Inches(0.65)
land.header_distance = Inches(0.25)
land.footer_distance = Inches(0.25)
heading(doc, "11. เปรียบเทียบ Cloud ต่างประเทศ 5 ราย", 1)
doc.add_paragraph("เกณฑ์เดียวกัน: x86-64, Public IPv4, Region Singapore/SEA, Ubuntu/Docker, 4 vCPU, RAM ≥8 GB, Disk ≥160 GB. แปลงที่ 33 บาท/USD; ไม่รวม VAT/ภาษี/โดเมน/egress เกินโควตา. ราคา ณ 26 ส.ค. 2026.")
global_rows = [
    ("1", "Hostinger KVM 4", "Singapore", "4 / 16 GB / 200 GB NVMe", "$28.99 Renewal\n($12.99 promo)", "Weekly รวม", "≈฿957", "แนะนำ: RAM สูงสุด/ราคา; สัญญา 2 ปี"),
    ("2", "Vultr VC2", "Singapore", "4 / 8 GB / 160 GB SSD", "$40", "+20% = $48", "฿1,320 / 1,584", "คล่องตัว; ราคา API โปร่งใส"),
    ("3", "DigitalOcean Basic", "SGP1", "4 / 8 GiB / 160 GiB", "$48", "+20% = $57.60", "฿1,584 / 1,901", "เอกสารดี; ใช้งานง่าย"),
    ("4", "Hetzner CPX32", "Singapore", "4 / 8 GB / 160 GB", "$57.99 + IPv4 $0.60", "+20% plan ≈ $70.19", "฿1,933 / 2,316", "Firewall ฟรี; traffic SG 0.5 TB"),
    ("5", "AWS Lightsail CO", "Singapore", "4 / 8 GB / 320 GB", "$84", "$0.05/GB-month", "฿2,772 + snapshot", "AWS ecosystem; ต้นทุนสูงกว่า"),
]
add_table(doc, ["อันดับ", "ผู้ให้บริการ", "Region", "vCPU/RAM/Disk", "Base/เดือน", "Backup/Snapshot", "งบบาท", "ข้อสังเกต"], global_rows, [0.55, 1.35, 0.85, 1.65, 1.25, 1.25, 1.15, 1.55], 7.8)
heading(doc, "11.1 คำแนะนำ", 2)
add_bullet(doc, "ใช้ KVM 2 ที่ติดตั้งแล้วเป็น Technical Pilot 7 วัน; หาก CPU/RAM/IOPS/พื้นที่หรือจำนวนผู้ใช้ไม่ผ่านเกณฑ์ ให้ resize เป็น KVM 4 ก่อน Production.")
add_bullet(doc, "หากต้องการย้าย/resize รายชั่วโมงและ UI ที่ตรงไปตรงมา ให้ Vultr หรือ DigitalOcean เป็นทางสำรอง.")
add_bullet(doc, "อย่าตัดสินใจจากราคาโปรโมชัน; งบควรใช้ Renewal/On-demand และแยก provider snapshot จาก DB backup ผ่าน SFTP.")

heading(doc, "11.2 คะแนนเชิงบริหาร (5 = ดีที่สุด)", 2)
add_table(doc, ["Provider", "Value", "Ease", "SEA", "Backup", "Enterprise", "สรุป"], [
    ("Hostinger", "5", "4", "5", "4", "3", "Best value"),
    ("Vultr", "4", "4", "5", "4", "3", "Flexible alternative"),
    ("DigitalOcean", "3", "5", "5", "4", "4", "Best developer experience"),
    ("Hetzner", "3", "4", "5", "4", "3", "Good controls; 2026 price increased"),
    ("AWS Lightsail", "2", "4", "5", "4", "5", "Best ecosystem"),
], [1.4, 0.8, 0.8, 0.8, 0.9, 1.0, 2.5], 8.4)
heading(doc, "11.3 วิธีอ่านคะแนน", 2)
add_bullet(doc, "Value ให้น้ำหนักต้นทุนฐานและทรัพยากรที่ได้; Ease ครอบคลุมการ provision, resize และเอกสาร; SEA พิจารณา region ใกล้ประเทศไทย.")
add_bullet(doc, "Backup ให้คะแนนความพร้อมของ snapshot/backup จากผู้ให้บริการ แต่ไม่ใช้แทน verified database backup ที่ดาวน์โหลดผ่าน SFTP.")
add_bullet(doc, "Enterprise สะท้อน ecosystem, governance และช่องทาง support. คะแนนนี้เป็น decision aid สำหรับ WorldFert ไม่ใช่ benchmark สากล.")
callout(doc, "ข้อสรุปการตัดสินใจ", "อนุมัติ Hostinger เป็น Pilot และกำหนด Vultr/DigitalOcean เป็นทางสำรอง หากผลวัด IOPS, latency หรือ support ไม่ผ่านเกณฑ์ก่อน Go-Live.", GREEN)

page_break(doc)
heading(doc, "12. ผู้ให้บริการในไทย 5 ราย", 1)
doc.add_paragraph("ราคา ReadyIDC/NIPA มาจากหน้าสาธารณะ; รายอื่นต้อง RFQ. ช่วงงบที่ระบุเป็น planning envelope ภายใน ไม่ใช่ใบเสนอราคาผู้ให้บริการ และต้องยืนยัน Public IPv4, Docker/MSSQL Linux, firewall ports และ backup policy ใน TOR.")
thai_rows = [
    ("ReadyIDC", "Xeon R3", "4 / 12 GB / 240 GB", "1 IPv4 + Daily backup 3 วัน", "฿1,000 ex VAT\n฿1,070 incl VAT", "ราคาสาธารณะคุ้มสุด; POC performance"),
    ("NIPA Cloud", "All Purpose + 200 GB", "4 / 16 GB / 200 GB", "External IP + Security Group", "≈฿4,740 ex VAT\n≈฿5,072 incl VAT", "Annual self-service; IIG extra"),
    ("INET", "Enterprise Cloud RFQ", "กำหนดตาม TOR", "VM + Backup + DRaaS + Container", "RFQ\nตั้งงบ 5–15k", "SLA 99.95; ทีมไทย/มาตรฐานองค์กร"),
    ("True IDC", "Cloud+ RFQ", "กำหนดตาม TOR", "Elastic IP + Pay-as-you-go", "RFQ\nตั้งงบ 5–15k", "Data residency ไทย; 24x7"),
    ("Cloud HM", "VMware IaaS RFQ", "กำหนดตาม TOR", "3 DC + all-flash + firewall", "RFQ\nตั้งงบ 7–20k", "Enterprise managed/SLA 99.9–99.95"),
]
add_table(doc, ["ผู้ให้บริการ", "แผน", "vCPU/RAM/Disk", "Public/Backup", "งบ/เดือน", "ข้อสังเกต"], thai_rows, [1.2, 1.45, 1.5, 1.7, 1.35, 2.1], 8.1)
heading(doc, "12.1 ข้อความที่ต้องใส่ใน RFQ/TOR", 2)
for item in [
    "VM x86-64 Ubuntu 24.04, root/sudo, อนุญาต Docker และ Microsoft SQL Server Linux container.",
    "Public IPv4 คงที่ ไม่มี CGNAT; inbound 22/80/443/1433/3306 และ Source CIDR firewall.",
    "4 vCPU dedicated/preferred, RAM 16 GB, NVMe/Premium SSD 200 GB, ขยาย disk online ได้.",
    "ราคาต่อเดือนรวม IPv4, backup/snapshot, domestic/international bandwidth, VAT และ support SLA.",
    "RPO/RTO, retention, restore fee, support 24x7, log/audit, data location และ exit/migration policy.",
]:
    add_bullet(doc, item)

# Back to portrait
portrait = doc.add_section(WD_SECTION.NEW_PAGE)
portrait.orientation = WD_ORIENT.PORTRAIT
portrait.page_width = Inches(8.5)
portrait.page_height = Inches(11)
portrait.top_margin = Inches(0.85)
portrait.bottom_margin = Inches(0.75)
portrait.left_margin = Inches(1)
portrait.right_margin = Inches(1)
portrait.header_distance = Inches(0.35)
portrait.footer_distance = Inches(0.35)
heading(doc, "ภาคผนวก A — Checklist ก่อนเริ่ม 4 ชั่วโมง", 1)
for c in [
    "[ ] C-Level/Business owner อนุมัติ provider, budget, downtime และ rollback authority",
    "[ ] VPS x86-64 พร้อม Public IPv4; Hostinger/provider firewall ถูก Apply และ UFW สอดคล้องหากเลือกเปิดใช้งาน",
    "[ ] DNS @/app/api/mssql/mysql/portainer/test/api-test ชี้ Public IPv4, www CNAME ถูกต้อง, TTL 300 และ MX/TXT เดิมยังอยู่",
    "[ ] Windows OpenSSH tools + deploy/SFTP keys ผ่าน",
    "[ ] .env ไม่มี CHANGE_ME และไม่ถูก commit",
    "[ ] Full backup ล่าสุด + SHA-256; ทดลอง upload/download ผ่าน",
    "[ ] Empty-stack deploy + health check ผ่านก่อน window",
    "[ ] ผู้ทำ UAT และช่องทาง War Room พร้อม",
    "[ ] เริ่ม Pilot Monitor 7 วันและกำหนดผู้รับแจ้งเตือน",
    "[ ] ระบบเดิมยังพร้อม rollback และไม่ถูกปิด/ลบ",
    "[ ] Change log ระบุเวลาทุก Gate และผู้อนุมัติ Go/No-Go",
]:
    doc.add_paragraph(c)

page_break(doc)
heading(doc, "ภาคผนวก B — ตำแหน่งไฟล์ส่งมอบ", 1)
add_table(doc, ["รายการ", "Path"], [
    ("คู่มือหลักฉบับเดียว", r"deliverables\cloud-deployment\WorldFert_Complete_Operations_Handbook_TH.docx"),
    ("ฉบับ PDF สำหรับอ่าน/พิมพ์", r"deliverables\cloud-deployment\WorldFert_Complete_Operations_Handbook_TH.pdf"),
    ("สไลด์ C-Level", r"deliverables\cloud-deployment\WorldFert_Cloud_Proposal_C-Level_TH.pptx"),
    ("Compose/README", "deploy\\cloud-vps\\"),
    ("DB client reference", r"deploy\cloud-vps\DB-CLIENT-CONNECTIONS.example.txt"),
    ("BAT scripts", "deploy\\cloud-vps\\windows\\"),
    ("Server scripts", "deploy\\cloud-vps\\server\\"),
    ("Credential vault (ห้ามแจก)", r"deploy\cloud-vps\.local-secrets\CREDENTIALS.txt"),
    ("Application admin (ห้ามแจก)", r"deploy\cloud-vps\.local-secrets\APPLICATION-ADMIN.txt"),
    ("Deploy/SFTP private keys (ห้ามแจก)", r"deploy\cloud-vps\.local-secrets\worldfert-hostinger-deploy และ worldfert-hostinger-sftp"),
    ("Database CA สำหรับ client", r"deploy\cloud-vps\.local-secrets\worldfert-db-ca.crt"),
], [2.0, 4.4], 9.3)
callout(doc, "แหล่งข้อมูลลับ", "รายการที่ระบุว่า ‘ห้ามแจก’ ต้องอยู่เฉพาะเครื่องผู้ดูแล/Password vault ที่องค์กรอนุมัติ. คู่มือฉบับนี้สามารถส่งให้ทีมปฏิบัติการได้เพราะไม่มีค่ารหัสผ่านหรือ private-key content.", AMBER)

page_break(doc)
heading(doc, "ภาคผนวก C — แหล่งอ้างอิงราคาและข้อจำกัด", 1)
sources = [
    ("Hostinger VPS pricing", "https://www.hostinger.com/vps-hosting?lang=en"),
    ("Hostinger server locations", "https://support.hostinger.com/en/articles/1583267-where-are-hostinger-servers-located"),
    ("Hostinger VPS backup", "https://support.hostinger.com/en/articles/1583232-how-to-back-up-or-restore-a-vps"),
    ("Vultr plans API", "https://api.vultr.com/v2/plans?type=vc2&per_page=100"),
    ("Vultr automatic backup pricing", "https://docs.vultr.com/support/platform/billing/how-much-does-it-cost-to-enable-automatic-backups"),
    ("DigitalOcean Droplet pricing", "https://www.digitalocean.com/pricing/droplets"),
    ("DigitalOcean backup pricing", "https://docs.digitalocean.com/products/backups/details/pricing/"),
    ("Hetzner 2026 Singapore pricing", "https://docs.hetzner.com/general/infrastructure-and-availability/price-adjustment/"),
    ("Hetzner Primary IPv4", "https://docs.hetzner.com/cloud/servers/primary-ips/overview/"),
    ("AWS Lightsail bundles", "https://docs.aws.amazon.com/lightsail/latest/userguide/amazon-lightsail-bundles.html"),
    ("AWS Lightsail snapshots", "https://docs.aws.amazon.com/lightsail/latest/userguide/amazon-lightsail-faq-snapshots.html"),
    ("ReadyIDC Cloud Server", "https://www.readyidc.com/en/products/cloud-server/"),
    ("NIPA Compute pricing", "https://nipa.cloud/th/pricing/nipa-space/compute-instance"),
    ("NIPA External IP pricing", "https://nipa.cloud/pricing/nipa-space/external-ip"),
    ("INET Cloud Solution", "https://www.inet.co.th/th/services/cloud-solution"),
    ("True IDC Cloud+", "https://www.trueidc.com/en/service/67/True-IDC-Cloud%2B"),
    ("Cloud HM infrastructure", "https://www.cloudhm.co.th/en/secured-infrastructure"),
    ("SQL Server 2022 edition limits", "https://learn.microsoft.com/en-us/sql/sql-server/editions-and-components-of-sql-server-2022"),
    ("SQL Server Linux container deployment", "https://learn.microsoft.com/en-us/sql/linux/quickstart-install-connect-docker"),
    ("MySQL require_secure_transport", "https://dev.mysql.com/doc/refman/8.0/en/server-system-variables.html"),
    ("Portainer CE on Docker", "https://docs.portainer.io/2.33-lts/start/install-ce/server/docker/linux"),
    ("Portainer behind reverse proxy", "https://docs.portainer.io/advanced/reverse-proxy"),
    ("Caddy automatic HTTPS", "https://caddyserver.com/docs/quick-starts/https"),
    ("SSMS Database Engine connection", "https://learn.microsoft.com/en-us/ssms/f1-help/connect-to-server-login-page-database-engine"),
    ("DBeaver SQL Server driver", "https://dbeaver.com/docs/dbeaver/Database-driver-Microsoft-SQL-Server/"),
    ("DBeaver SSL configuration", "https://dbeaver.com/docs/dbeaver/SSL-Configuration/"),
    ("Bank of Thailand FX rates", "https://www.bot.or.th/en/statistics/exchange-rate.html"),
]
for label, url in sources:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    add_hyperlink(p, label, url)
    spacer = p.add_run(" ")
    spacer.font.size = Pt(10)

doc.core_properties.title = "WorldFert Complete Cloud VPS Operations Handbook"
doc.core_properties.subject = "Single operational handbook: VPS creation, credentials registry, Docker deploy, database restore, access, testing, backup and maintenance"
doc.core_properties.author = "WorldFert Project Team"
doc.core_properties.keywords = "WorldFert, Hostinger, Docker Compose, Portainer, SFTP, MSSQL, MySQL, Cloud VPS, Operations Handbook"

doc.save(OUT)
print(OUT)
